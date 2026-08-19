# -*- coding: utf-8 -*-
"""
Generator Laporan Pelaksanaan Tugas Tenaga Ahli (7 anggota) — BULAN KE-2.
Kegiatan: Penyediaan dan Pengelolaan Data Statistik Pariwisata Jakarta — TA 2026.

Fokus bulan ke-2:
  1) Platform data (dispar.rantai.dev) berjalan mandiri (self-host, DB sendiri).
  2) Data BARU hasil crawl: wisata ramah muslim/halal, wisatawan mancanegara,
     kuliner, seni & pertunjukan (indikator GCI Cultural Experience).
  3) Data yang DIPERKUAT: geocoding dataset "Jumlah Pengunjung Event 2026"
     (lokasi -> alamat/lat/lon/sumber) + verifikasi ulang 12 event prioritas.

Menghasilkan .docx (python-docx). Konversi ke .pdf via LibreOffice (langkah terpisah).
"""
import csv, json, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

BASE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(BASE, "assets-bulan2")
REPO = os.path.abspath(os.path.join(BASE, "..", ".."))
DATA = os.path.join(REPO, "platform", "data")
INK = RGBColor(0x1a, 0x1a, 0x1a)
GREY = RGBColor(0x55, 0x55, 0x55)

# ---------- helpers ----------

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table, color="7f7f7f", sz="4"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)

def cell_text(cell, text, bold=False, size=10, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align: p.alignment = align
    run = p.add_run(text)
    run.bold = bold; run.font.size = Pt(size); run.font.name = "Arial"; run.font.color.rgb = color
    return p

def para(doc, text, size=11, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         color=INK, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after); pf.space_before = Pt(space_before); pf.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    run.font.name = "Arial"; run.font.color.rgb = color
    return p

def heading(doc, text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(size); run.font.name = "Arial"; run.font.color.rgb = INK
    return p

def bullets(doc, items, letters=False):
    for i, it in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.9); p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        prefix = f"{chr(97+i)}.  " if letters else "•  "
        r = p.add_run(prefix); r.font.name="Arial"; r.font.size=Pt(11); r.bold = letters
        r2 = p.add_run(it); r2.font.name="Arial"; r2.font.size=Pt(11); r2.font.color.rgb = INK

def add_image(doc, filename, caption, width_cm=15.5, max_h_cm=19.0):
    path = os.path.join(ASSETS, filename)
    # skala proporsional; batasi tinggi agar tidak melebihi satu halaman
    w_px, h_px = Image.open(path).size
    w = width_cm
    h = width_cm * h_px / w_px
    if h > max_h_cm:
        h = max_h_cm; w = max_h_cm * w_px / h_px
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Cm(w), height=Cm(h))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9)
    r.font.name = "Arial"; r.font.color.rgb = GREY

def realisasi_table(doc, rows):
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_table_borders(t)
    hdr = t.rows[0].cells
    cell_text(hdr[0], "Uraian Tugas (KAK)", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_text(hdr[1], "Realisasi Bulan Ini", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    for c in hdr: set_cell_bg(c, "d9d9d9")
    for tugas, real in rows:
        r = t.add_row().cells
        cell_text(r[0], tugas, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
        cell_text(r[1], real, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    t.columns[0].width = Cm(6.6); t.columns[1].width = Cm(9.0)
    for row in t.rows:
        row.cells[0].width = Cm(6.6); row.cells[1].width = Cm(9.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def data_table(doc, header, rows, widths):
    t = doc.add_table(rows=1, cols=len(header)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; set_table_borders(t)
    for i, h in enumerate(header):
        cell_text(t.rows[0].cells[i], h, bold=True, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_bg(t.rows[0].cells[i], "d9d9d9")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cell_text(cells[i], v, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ---------- data samples from real deliverables (this month) ----------

def json_rows(fname):
    d = json.load(open(os.path.join(DATA, fname), encoding="utf-8"))
    return d.get("rows", d) if isinstance(d, dict) else d

def clip(s, n):
    s = "" if s is None else str(s)
    return s if len(s) <= n else s[: n - 1] + "…"

# Halal — restoran & zona KHAS tersertifikasi (crawl BPJPH/LPPOM MUI)
_hal = json_rows("halal-restoran-halal.json")
HAL_HEADER = ["No.", "Nama", "Jenis", "Kota", "No. Sertifikat Halal", "Lembaga"]
HAL_ROWS = [[str(i+1), clip(r["nama"],22), clip(r["jenis"],20), clip(r["kota"],14),
             clip(r["no_sertifikat_halal"],22), clip(r["lembaga_sertifikasi"],14)]
            for i, r in enumerate(_hal[:6])]
HAL_WIDTHS = [0.8, 3.6, 3.2, 2.3, 3.6, 2.0]

# Event — dataset diperkuat (geocode: lokasi -> alamat/lat/lon)
_evt = json_rows("event-visitors-2026.json")
_want = ["BTN JAKARTA INTERNATIONAL MARATHON (BTN JAKIM) 2026", "RED CROSS INDONESIA - PMI",
         "TRAVEL MEET ASIA 2026", "JAKARTA AT IMEX FRANKFURT 2026",
         "SALES MISSION CHINA AT XIAMEN 2026", "JAKARTA AT ITB CHINA 2026"]
_evt_by = {r["nama_event"]: r for r in _evt}
EVT_HEADER = ["Nama Event", "Alamat (hasil geocode)", "Kota", "Lat", "Lon"]
EVT_ROWS = []
for nm in _want:
    r = _evt_by.get(nm)
    if r:
        EVT_ROWS.append([clip(r["nama_event"],26), clip(r["alamat"],40), clip(r["kota"],16),
                         clip(r["lat"],9), clip(r["lon"],9)])
EVT_WIDTHS = [4.0, 5.6, 2.3, 1.8, 1.8]

# Wisman — agregasi total kunjungan per negara (top 6)
_wis = json_rows("wisman-per-negara.json")
_agg = {}
for r in _wis:
    try: v = int(float(r["jumlah_kunjungan"]))
    except (ValueError, TypeError): v = 0
    _agg[r["negara"]] = _agg.get(r["negara"], 0) + v
_top = sorted(_agg.items(), key=lambda kv: kv[1], reverse=True)[:6]
WIS_HEADER = ["No.", "Negara Asal", "Total Kunjungan (Sem. I 2026)"]
WIS_ROWS = [[str(i+1), neg, f"{tot:,}".replace(",", ".")] for i, (neg, tot) in enumerate(_top)]
WIS_WIDTHS = [1.0, 6.0, 6.0]

# ---------- document skeleton ----------

def build(doc_meta):
    doc = Document()
    st = doc.styles['Normal']; st.font.name = "Arial"; st.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.2); section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.2)

    # KOP
    para(doc, "PEMERINTAH PROVINSI DAERAH KHUSUS IBUKOTA JAKARTA", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, "DINAS PARIWISATA DAN EKONOMI KREATIF", size=13, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, "Bidang Data, Informasi dan Pengembangan Destinasi", size=10, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=2)
    hr = doc.add_paragraph(); hr.paragraph_format.space_after = Pt(10)
    pPr = hr._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'12')
    bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),'1a1a1a'); pbdr.append(bottom); pPr.append(pbdr)

    para(doc, "LAPORAN PELAKSANAAN TUGAS TENAGA AHLI", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    para(doc, "Kegiatan Penyediaan dan Pengelolaan Data Statistik Pariwisata Jakarta — Tahun Anggaran 2026",
         size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=10)

    # IDENTITAS
    t = doc.add_table(rows=0, cols=2); set_table_borders(t, color="bfbfbf")
    ident = [
        ("Nama Tenaga Ahli", "[Nama Tenaga Ahli]"),
        ("NIK / No. Kontrak", "[NIK / No. Kontrak]"),
        ("Jabatan / Posisi", doc_meta["jabatan"]),
        ("Kualifikasi", doc_meta["kualifikasi"]),
        ("Periode Laporan", "Bulan ke-2 (Laporan Bulanan Kedua)"),
        ("Sub Kegiatan", "Perencanaan Daya Tarik Wisata Provinsi"),
        ("Lokasi", "Provinsi DKI Jakarta"),
    ]
    for k, v in ident:
        row = t.add_row().cells
        cell_text(row[0], k, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT); set_cell_bg(row[0], "eeeeee")
        cell_text(row[1], v, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        row[0].width = Cm(4.8); row[1].width = Cm(10.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    heading(doc, "I.  PENDAHULUAN")
    para(doc, PENDAHULUAN_UMUM)
    para(doc, doc_meta["pendahuluan_peran"])

    heading(doc, "II.  URAIAN TUGAS")
    para(doc, "Sesuai Kerangka Acuan Kerja (KAK) kegiatan, uraian tugas untuk posisi ini adalah sebagai berikut:",
         space_after=4)
    bullets(doc, doc_meta["uraian_tugas"], letters=True)

    heading(doc, "III.  PELAKSANAAN PEKERJAAN BULAN INI")
    para(doc, doc_meta["pelaksanaan_intro"])
    realisasi_table(doc, doc_meta["realisasi"])

    heading(doc, "IV.  BUKTI PELAKSANAAN")
    para(doc, doc_meta["bukti_intro"], space_after=6)
    for item in doc_meta["bukti"]:
        if item[0] == "img":
            add_image(doc, item[1], item[2])
        elif item[0] == "hal":
            para(doc, item[1], bold=True, size=10.5, space_before=4, space_after=3,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
            data_table(doc, HAL_HEADER, HAL_ROWS, HAL_WIDTHS)
            para(doc, "Sumber: dataset Restoran & Zona KHAS Tersertifikasi Halal Jakarta (134 baris; crawl BPJPH/LPPOM MUI + geocode OSM/Nominatim).",
                 size=8.5, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
        elif item[0] == "evt":
            para(doc, item[1], bold=True, size=10.5, space_before=4, space_after=3,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
            data_table(doc, EVT_HEADER, EVT_ROWS, EVT_WIDTHS)
            para(doc, "Sumber: dataset Jumlah Pengunjung Event Jakarta 2026 (804 baris); kolom lokasi diperkuat menjadi alamat, koordinat, dan sumber via geocoding + penelusuran web.",
                 size=8.5, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
        elif item[0] == "wis":
            para(doc, item[1], bold=True, size=10.5, space_before=4, space_after=3,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
            data_table(doc, WIS_HEADER, WIS_ROWS, WIS_WIDTHS)
            para(doc, "Sumber: dataset Wisatawan Mancanegara per Negara Asal (agregat Semester I 2026).",
                 size=8.5, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)

    heading(doc, "V.  RENCANA TINDAK LANJUT — PENGEMBANGAN PLATFORM AI-DATA")
    para(doc, RTL_UMUM)
    para(doc, doc_meta["rtl_peran"])
    for item in doc_meta.get("rtl_bukti", []):
        add_image(doc, item[0], item[1])

    heading(doc, "VI.  PENUTUP")
    para(doc, doc_meta["penutup"])

    # tanda tangan
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    sig = doc.add_table(rows=0, cols=2)
    r = sig.add_row().cells
    cell_text(r[0], "Mengetahui,\nPejabat Pembuat Komitmen", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(r[1], "Jakarta, ......................... 2026\nTenaga Ahli yang bersangkutan",
              size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4): sig.add_row()
    r = sig.add_row().cells
    cell_text(r[0], "(  Bima Agung  )", bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(r[1], "(  [Nama Tenaga Ahli]  )", bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = sig.add_row().cells
    cell_text(r[0], "NIP. 197907162011011008", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
    cell_text(r[1], doc_meta["jabatan"], size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)

    outdir = os.path.join(BASE, "bulan-2")
    os.makedirs(outdir, exist_ok=True)
    out = os.path.join(outdir, doc_meta["filename"] + ".docx")
    doc.save(out)
    print("wrote", out)

# ---------- shared narrative ----------

PENDAHULUAN_UMUM = (
    "Laporan ini merupakan laporan bulanan kedua atas pelaksanaan kegiatan Penyediaan dan Pengelolaan "
    "Data Statistik Pariwisata Jakarta Tahun Anggaran 2026. Bila bulan pertama difokuskan pada pembangunan "
    "fondasi data indikator awal dan dashboard percontohan, maka bulan kedua difokuskan pada tiga hal: "
    "(1) mengoperasikan platform data pariwisata secara mandiri melalui infrastruktur milik Dinas "
    "(dapat diakses pada dispar.rantai.dev), lepas dari ketergantungan layanan pihak ketiga; "
    "(2) memperluas cakupan data melalui pengumpulan (crawl) data baru — wisata ramah muslim/halal, "
    "wisatawan mancanegara, kuliner, serta seni dan pertunjukan sebagai pendukung indikator Global City "
    "Index (GCI) dimensi Pengalaman Budaya; dan (3) memperkuat kualitas data yang sudah ada, khususnya "
    "pengayaan (enrichment) lokasi event menjadi alamat, titik koordinat, dan sumber yang terverifikasi."
)

RTL_UMUM = (
    "Sebagai kelanjutan, tim melanjutkan pengembangan Jakarta Tourism Intelligence Platform — satu fondasi "
    "AI-Data yang menggabungkan lima kemampuan inti: Data Gathering, Lakehouse Making, Dashboarding, AI "
    "Analyst, dan Policy Insight Layer. Pada bulan kedua ini, kemampuan Dashboarding dan Data Gathering telah "
    "mulai terwujud nyata dan berjalan di atas infrastruktur milik Dinas; tahap berikutnya menyiapkan lapisan "
    "AI Analyst agar pengguna dapat bertanya langsung kepada data dengan jawaban yang tetap dapat diaudit."
)

# ---------- per-member content ----------

MEMBERS = []

# 1. Project Manager
MEMBERS.append(dict(
    filename="01-PM-Analisis-Data-Dashboard",
    jabatan="Project Manager Analisis Data dan Pembangunan Dashboard Pariwisata",
    kualifikasi="S2 / setara, pengalaman minimal 5 tahun",
    pendahuluan_peran=(
        "Sebagai Project Manager Analisis Data dan Pembangunan Dashboard Pariwisata, tenaga ahli bertanggung "
        "jawab merencanakan dan mengoordinasikan seluruh sumber daya tim, menetapkan prioritas pekerjaan bulan "
        "kedua (operasionalisasi platform mandiri, perluasan data, dan penguatan mutu data), memantau progres, "
        "mengelola risiko, serta memastikan seluruh keluaran selesai sesuai standar kualitas dan harapan pemilik program."
    ),
    uraian_tugas=[
        "Mengelola dan merencanakan SDM yang dibutuhkan untuk kegiatan Analisis Data dan Pembangunan Dashboard Pariwisata, termasuk analisis, desain, pelaksanaan, pengujian/validasi, dan output yang dihasilkan;",
        "Memberikan laporan kegiatan Analisis Data dan Pembangunan Dashboard Pariwisata berupa progres pekerjaan dan pembaruan;",
        "Memastikan kegiatan selesai tepat waktu, sesuai anggaran, memenuhi standar kualitas, dan harapan pemilik program;",
        "Mengelola risiko dan isu pekerjaan, memastikan penyediaan informasi tepat waktu, serta melakukan mitigasi risiko dan langkah eskalasi.",
    ],
    pelaksanaan_intro=(
        "Pada bulan kedua, tenaga ahli memimpin operasionalisasi platform data pariwisata secara mandiri pada "
        "infrastruktur milik Dinas (dispar.rantai.dev), mengoordinasikan perluasan katalog data hingga menembus "
        "182 dataset, serta mengarahkan penambahan tiga dashboard indikator GCI Pengalaman Budaya (kuliner, seni "
        "& pertunjukan, wisatawan mancanegara). Rincian realisasi terhadap uraian tugas sebagai berikut:"
    ),
    realisasi=[
        ("Perencanaan SDM & prioritas pekerjaan",
         "Menetapkan tiga prioritas bulan kedua (platform mandiri, perluasan data baru, penguatan mutu data) dan membagi peran tim antara pengembangan dashboard, pipeline crawl data baru, dan geocoding/verifikasi."),
        ("Pelaporan progres & pembaruan",
         "Menyusun rekapitulasi capaian: platform berjalan mandiri dengan 182 dataset terkatalog; penambahan 6 dataset wisata ramah muslim/halal, dataset wisatawan mancanegara (per bulan/negara/pintu masuk), serta penguatan 804 baris data pengunjung event via geocoding."),
        ("Ketepatan waktu, mutu & anggaran",
         "Memastikan pemindahan platform ke infrastruktur mandiri berjalan tanpa biaya langganan basis data pihak ketiga, sekaligus menekan biaya enumerasi data dengan mengutamakan sumber terbuka (OSM/Nominatim) dan penelusuran web bersumber."),
        ("Manajemen risiko & isu",
         "Mengidentifikasi risiko kedaulatan & keberlanjutan data (ketergantungan layanan pihak ketiga) dan menetapkan mitigasi berupa self-host di server Dinas; mengidentifikasi risiko akurasi lokasi event dan menetapkan mitigasi verifikasi ulang bertingkat (confidence)."),
    ],
    bukti_intro=(
        "Berikut bukti keluaran kegiatan di bawah koordinasi Project Manager: platform data yang berjalan mandiri "
        "beserta katalog dan indikator GCI Pengalaman Budaya."
    ),
    bukti=[
        ("img", "01-home.png", "Gambar 1. Beranda platform data pariwisata & ekraf DKI Jakarta (dispar.rantai.dev) — berjalan mandiri pada infrastruktur Dinas."),
        ("img", "06-sdi-katalog.png", "Gambar 2. Katalog data (Satu Data) — 182 dataset terkatalog, terbagi tier Primer & Sekunder."),
        ("img", "02-gci.png", "Gambar 3. Dashboard GCI Jakarta — dimensi Pengalaman Budaya dengan empat sub-indikator pariwisata."),
    ],
    rtl_peran=(
        "Dalam pengembangan platform, Project Manager berperan memutakhirkan roadmap: Fase 1 (fondasi data & "
        "dashboard) telah terpenuhi; bulan ini memasuki Fase 2 (platform data mandiri, perluasan indikator & "
        "sumber data); selanjutnya menyiapkan AI Analyst dan Policy Insight Layer pada Fase 3 — dengan setiap "
        "fase tetap selaras dengan KPI Disparekraf dan indikator daya saing kota global."
    ),
    rtl_bukti=[],
    penutup=(
        "Secara keseluruhan, pekerjaan bulan kedua telah menempatkan platform data pada infrastruktur mandiri, "
        "memperluas cakupan data, dan memperkuat mutu data yang sudah ada. Koordinasi tim dilanjutkan untuk "
        "menambah indikator dan menyiapkan lapisan AI Analyst pada tahap berikutnya."
    ),
))

# 2. Database Administrator
MEMBERS.append(dict(
    filename="02-Database-Administrator",
    jabatan="Tenaga Ahli Database Administrator (Senior)",
    kualifikasi="S1 Teknik Informatika/Ilmu Komputer/Sistem Informasi, pengalaman minimal 7 tahun",
    pendahuluan_peran=(
        "Sebagai Database Administrator, tenaga ahli bertanggung jawab menyiapkan dan mengelola basis data "
        "statistik pariwisata pada infrastruktur mandiri Dinas, menyusun skema tabel untuk dataset baru, "
        "melaksanakan loading data hasil pengumpulan dan pengayaan, serta menjaga integritas, keterlacakan, "
        "dan keberlanjutan basis data."
    ),
    uraian_tugas=[
        "Merancang struktur basis data statistik pariwisata pada platform Database (logical & physical data model);",
        "Menyusun skema tabel, indeks, constraint, view, dan partisi sesuai kebutuhan analitik;",
        "Melaksanakan loading data hasil pengumpulan, pembersihan, dan validasi ke dalam Database;",
        "Melakukan optimasi kinerja basis data (query tuning, indexing, statistik, partitioning);",
        "Menyusun tata kelola akses dan keamanan basis data (role, privilege, audit trail);",
        "Menyusun pedoman backup, recovery, dan pemeliharaan basis data;",
        "Mendokumentasikan basis data (ERD, kamus tabel, panduan operasional).",
    ],
    pelaksanaan_intro=(
        "Pada bulan kedua, tenaga ahli memindahkan basis data platform ke instans PostgreSQL milik Dinas "
        "(self-host, terkontainerisasi) sehingga lepas dari layanan basis data pihak ketiga, serta menambahkan "
        "skema untuk dataset baru (wisata ramah muslim/halal, wisatawan mancanegara, event) dan memuatnya melalui "
        "skrip seed yang idempoten. Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Basis data mandiri & terkontainerisasi",
         "Menyiapkan instans PostgreSQL milik Dinas melalui kontainer (Docker/Portainer) dengan volume data persisten, sehingga platform berjalan mandiri dan datanya berada di lingkungan Dinas."),
        ("Skema tabel dataset baru",
         "Menetapkan skema baris-generik (record + dataset_column + dataset_sync) untuk memuat dataset baru: 6 dataset halal/ramah muslim, wisman (per bulan/negara/pintu masuk), TripAdvisor kuliner, dan indikator GCI/GPCI."),
        ("Loading data terkurasi & tervalidasi",
         "Memuat dataset baru dan dataset yang diperkuat ke basis data melalui skrip seed idempoten (hapus-lalu-muat) sehingga pemutakhiran data dapat diulang tanpa duplikasi; a.l. 804 baris event, 134 baris restoran halal, ratusan baris wisman."),
        ("Keterlacakan & integritas",
         "Menyimpan kolom sumber, nomor sertifikat halal (BPJPH/LPPOM MUI), serta koordinat pada setiap baris sebagai jejak keterlacakan; menjaga konsistensi tipe kolom antar-dataset."),
        ("Pemeliharaan & keberlanjutan",
         "Menerapkan mekanisme sinkronisasi terjadwal untuk dataset Satu Data Jakarta dan pemutakhiran manual terkendali untuk dataset sekunder, dengan volume basis data yang dicadangkan."),
        ("Dokumentasi struktur data",
         "Menyusun kamus kolom untuk dataset baru (halal, wisman, event) yang menyertai berkas deliverable dan halaman detail katalog."),
    ],
    bukti_intro=(
        "Bukti berupa struktur tabel dataset baru yang tampil pada halaman detail katalog beserta atribut "
        "keterlacakan (sumber, nomor sertifikat, koordinat)."
    ),
    bukti=[
        ("img", "08-halal-restoran.png", "Gambar 1. Struktur tabel dataset restoran halal — kolom terstandar Nama, Jenis, Alamat, Kota, Lat, Lon, No. Sertifikat Halal, Lembaga Sertifikasi, Sumber."),
        ("hal", "Cuplikan skema & isi tabel restoran halal (deliverable):"),
        ("img", "07-event-geocode.png", "Gambar 2. Halaman detail dataset event (804 baris) dengan kolom hasil pengayaan (alamat, koordinat, sumber) dan opsi unduh CSV/XLSX."),
    ],
    rtl_peran=(
        "Pada pengembangan platform, DBA berperan memantapkan lapisan Lakehouse (Raw–Clean–Business Data Mart) di "
        "atas basis data mandiri Dinas, lengkap dengan metadata dan jejak sumber, sebagai fondasi yang tervalidasi "
        "dan siap dipakai ulang oleh dashboard maupun AI Analyst."
    ),
    rtl_bukti=[],
    penutup=(
        "Basis data platform telah berjalan mandiri di lingkungan Dinas, memuat dataset baru dan dataset yang "
        "diperkuat secara terlacak. Tahap berikutnya adalah memformalkan pedoman backup/recovery dan memperluas "
        "skema untuk entitas pariwisata lain."
    ),
))

# 3. MDM Specialist
MEMBERS.append(dict(
    filename="03-MDM-Specialist",
    jabatan="Tenaga Ahli Master Data Management (MDM) Specialist (Senior)",
    kualifikasi="S1 Sistem Informasi/Teknik Informatika/Ilmu Komputer/Statistik, pengalaman minimal 7 tahun",
    pendahuluan_peran=(
        "Sebagai MDM Specialist, tenaga ahli bertanggung jawab membentuk dan menjaga data induk (golden record) "
        "entitas pariwisata dari data baru maupun data yang diperkuat, melalui konsolidasi, deduplikasi, verifikasi, "
        "dan validasi lintas sumber agar setiap entitas tepercaya dan dapat diaudit."
    ),
    uraian_tugas=[
        "Mengidentifikasi entitas master data pariwisata Jakarta dan menyusun struktur data master;",
        "Menyusun kebijakan dan prosedur MDM (data ownership, data stewardship, alur perubahan, persetujuan, dan audit);",
        "Melaksanakan pembentukan master data melalui konsolidasi, deduplikasi, verifikasi, dan validasi data dari berbagai sumber;",
        "Menyusun mekanisme pemeliharaan dan pemutakhiran master data secara berkala;",
        "Memastikan integrasi master data ke dalam database dan ketersediaannya bagi Dashboard melalui metadata.",
    ],
    pelaksanaan_intro=(
        "Pada bulan kedua, pekerjaan MDM yang menonjol adalah (1) pembentukan master data wisata ramah muslim/halal "
        "dengan nomor sertifikat halal sebagai kunci identitas tepercaya, dan (2) penguatan golden record event "
        "melalui pengayaan lokasi menjadi alamat/koordinat/sumber beserta verifikasi ulang 12 event prioritas. "
        "Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Identifikasi entitas & struktur master",
         "Menetapkan entitas master baru: restoran/zona KHAS halal, hotel ramah muslim, mall ramah muslim, RPH halal, produk kreatif halal, serta entitas event; menyusun atribut kunci masing-masing."),
        ("Kebijakan & prosedur stewardship",
         "Menetapkan aturan kepemilikan nilai per-sumber; nilai hasil riset tidak difabrikasi — bila tak terverifikasi, atribut dikembalikan ke status kosong atau ditandai untuk ditinjau (needs_review)."),
        ("Konsolidasi, deduplikasi & verifikasi",
         "Memverifikasi nomor sertifikat halal (BPJPH/LPPOM MUI) sebagai kunci induk; menyelaraskan penamaan dan menautkan tautan Google Maps yang dapat diverifikasi pada entitas halal."),
        ("Penguatan golden record event",
         "Memperkaya 804 baris event dengan alamat, titik koordinat, dan sumber; melakukan verifikasi ulang bertingkat pada 12 event prioritas hingga 7 mencapai keyakinan tinggi (high), termasuk koreksi alamat luar negeri yang keliru pada sumber asli."),
        ("Integrasi master ke dashboard",
         "Memastikan atribut master (nama, alamat, koordinat, sumber, nomor sertifikat, tingkat keyakinan) tersedia konsisten pada halaman detail katalog dan peta."),
    ],
    bukti_intro=(
        "Bukti berupa data induk halal dengan nomor sertifikat sebagai kunci identitas, dan golden record event "
        "yang telah diperkuat dengan alamat & koordinat."
    ),
    bukti=[
        ("img", "08-halal-restoran.png", "Gambar 1. Data induk restoran halal dengan No. Sertifikat Halal (BPJPH/LPPOM MUI) sebagai kunci identitas tepercaya dan sumber yang dapat diverifikasi."),
        ("evt", "Cuplikan golden record event yang diperkuat (12 event prioritas — hasil verifikasi ulang):"),
    ],
    rtl_peran=(
        "Pada pengembangan platform, MDM Specialist berperan memastikan setiap sumber yang masuk melalui Data "
        "Gathering dikonsolidasikan menjadi master data tepercaya — dengan kunci identitas yang jelas (mis. nomor "
        "sertifikat halal) dan tingkat keyakinan yang tercatat — sebelum dipakai dashboard maupun AI Analyst."
    ),
    rtl_bukti=[],
    penutup=(
        "Master data entitas halal telah dibentuk dengan kunci sertifikat, dan golden record event telah diperkuat "
        "serta diverifikasi ulang. Tahap berikutnya adalah memformalkan kebijakan stewardship dan memperluas master "
        "data ke seluruh entitas pariwisata."
    ),
))

# 4. BI Developer
MEMBERS.append(dict(
    filename="04-BI-Developer-1",
    jabatan="Tenaga Ahli BI Developer 1 (Junior)",
    kualifikasi="S1 Sistem Informasi/Teknik Informatika/Ilmu Komputer/Statistik, pengalaman minimal 5 tahun",
    pendahuluan_peran=(
        "Sebagai BI Developer, tenaga ahli bertanggung jawab membangun dashboard interaktif indikator pariwisata, "
        "menyusun visualisasi yang konsisten dengan definisi data, serta menyediakan fitur peta, penyaringan, dan "
        "ekspor bagi pengguna di lingkungan Dinas."
    ),
    uraian_tugas=[
        "Mengkonfigurasi koneksi Dashboard ke database melalui semantic layer/metadata MDM;",
        "Membangun dashboard interaktif (overview pariwisata, kunjungan wisatawan, akomodasi, destinasi, kontribusi ekonomi, dan dashboard tematik lainnya);",
        "Menyusun visualisasi yang konsisten dengan standar metadata dan definisi data;",
        "Mengatur akses pengguna, jadwal refresh data, dan pemeliharaan dashboard;",
        "Menyusun panduan pemanfaatan dashboard bagi pengguna di lingkungan Dinas.",
    ],
    pelaksanaan_intro=(
        "Pada bulan kedua, tenaga ahli membangun tiga dashboard indikator GCI dimensi Pengalaman Budaya — Kuliner "
        "(Michelin), Seni & Pertunjukan (lengkap peta venue), dan Wisatawan Mancanegara — beserta ringkasan, grafik, "
        "peta, filter, dan ekspor data. Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Koneksi dashboard ke sumber data",
         "Menghubungkan dashboard ke basis data mandiri melalui lapisan metadata; setiap kartu/grafik menautkan sumber ke katalog dataset agar dapat ditelusuri."),
        ("Dashboard interaktif tematik",
         "Membangun dashboard Kuliner (Michelin), Seni & Pertunjukan, dan Wisatawan Mancanegara pada dimensi Pengalaman Budaya, lengkap kartu ringkasan, tren bulanan, komposisi negara, dan peringkat kecamatan/venue."),
        ("Peta venue & visualisasi geospasial",
         "Menambahkan peta interaktif (Leaflet) venue seni & pertunjukan berbasis titik koordinat hasil geocoding, sehingga sebaran lokasi dapat dibaca langsung pada peta."),
        ("Visualisasi konsisten dengan definisi",
         "Menyelaraskan angka dashboard dengan definisi metadata (mis. penanda sumber & periode data); memakai pustaka grafik yang seragam untuk keterbacaan."),
        ("Filter, pencarian & ekspor",
         "Menyediakan filter tahun/kategori/wilayah, pencarian, serta tombol unduh CSV dan XLSX pada setiap dataset agar mudah dimanfaatkan pengguna Dinas."),
    ],
    bukti_intro=(
        "Bukti berupa tangkapan layar tiga dashboard indikator Pengalaman Budaya yang dibangun bulan ini."
    ),
    bukti=[
        ("img", "04-kuliner.png", "Gambar 1. Dashboard Kuliner (Michelin) — ringkasan, restoran per wilayah, Top-10 kecamatan, dan komposisi wilayah."),
        ("img", "03-seni-venuemap.png", "Gambar 2. Dashboard Seni & Pertunjukan — ringkasan, tren event/bulan, Top-10 tempat, dan peta venue interaktif (Leaflet)."),
        ("img", "05-wisman.png", "Gambar 3. Dashboard Wisatawan Mancanegara — 2,76 juta kunjungan, tren bulanan, Top-10 negara, dan komposisi negara."),
    ],
    rtl_peran=(
        "Pada pengembangan platform, BI Developer berperan memperluas lapisan Dashboarding menjadi indikator "
        "strategis penuh (destinasi, event & MICE, okupansi hotel, sentimen, efektivitas promosi) serta menyiapkan "
        "integrasi dengan AI Analyst agar pengguna dapat bertanya langsung kepada data."
    ),
    rtl_bukti=[
        ("02-gci.png", "Gambar 4. Dimensi Pengalaman Budaya sebagai kerangka penataan dashboard indikator GCI pariwisata."),
    ],
    penutup=(
        "Tiga dashboard indikator GCI Pengalaman Budaya telah dibangun, dapat diakses, dan dilengkapi peta serta "
        "ekspor data. Tahap berikutnya adalah menambah dashboard tematik dan menautkannya dengan lapisan AI Analyst."
    ),
))

# 5. Data Engineer
MEMBERS.append(dict(
    filename="05-Data-Engineer-2",
    jabatan="Tenaga Ahli Data Engineer 2 (Junior)",
    kualifikasi="S1 Teknik Informatika/Ilmu Komputer/Sistem Informasi, pengalaman minimal 6 tahun",
    pendahuluan_peran=(
        "Sebagai Data Engineer, tenaga ahli bertanggung jawab merancang dan menjalankan pipeline pengumpulan "
        "(crawl) dan integrasi data dari berbagai sumber, membersihkan dan menstandarkan raw data, serta "
        "menyiapkan pipeline geocoding dan penyajian platform pada infrastruktur mandiri."
    ),
    uraian_tugas=[
        "Merancang dan melaksanakan pipeline pengumpulan dan integrasi data (ETL/ELT) dari berbagai sumber ke dalam Database;",
        "Membersihkan, mentransformasi, dan menstandarisasi raw data agar siap digunakan untuk analisis dan visualisasi;",
        "Memastikan integrasi data multi-sumber dilakukan dengan menjaga kualitas dan keterlacakan data;",
        "Menerapkan standar metadata dan master data dalam proses integrasi data;",
        "Menjamin keamanan data dan kepatuhan terhadap tata kelola data pada proses integrasi.",
    ],
    pelaksanaan_intro=(
        "Pada bulan kedua, tenaga ahli membangun sejumlah pipeline crawl data baru (wisata ramah muslim/halal, "
        "wisatawan mancanegara, kuliner, artis/pertunjukan), pipeline geocoding untuk memperkuat data event, serta "
        "menyiapkan pipeline penyajian (kontainerisasi & tunnel) agar platform dapat diakses publik dari server "
        "Dinas. Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Pipeline pengumpulan (crawl) data baru",
         "Membangun pipeline crawl data wisata ramah muslim/halal dari sumber resmi & media (BPJPH/LPPOM MUI, situs resmi) dan pipeline data wisatawan mancanegara (per bulan, per negara asal, per pintu masuk) beserta data kuliner dan artis Top Chart."),
        ("Pipeline geocoding & pengayaan",
         "Membangun pipeline geocoding (OpenStreetMap/Nominatim + penelusuran web bersumber) yang mengubah teks lokasi menjadi alamat, titik koordinat (lat/lon), dan sumber pada 804 baris event serta entitas halal."),
        ("Pembersihan & standardisasi",
         "Menormalkan penulisan tanggal, angka, alamat, dan penamaan; menyaring wilayah non-DKI; menstandarkan kolom antar-dataset agar konsisten saat diintegrasikan."),
        ("Integrasi multi-sumber & keterlacakan",
         "Menggabungkan data lintas sumber dengan menandai asal setiap nilai (sumber, nomor sertifikat, tingkat keyakinan) untuk menjaga keterlacakan."),
        ("Pipeline penyajian mandiri & efisiensi",
         "Mengontainerisasi aplikasi (Docker) dan menyiapkan akses publik melalui tunnel aman ke server Dinas; mengutamakan sumber gratis dan membatasi pemanggilan API berbayar agar tetap efisien dan patuh."),
    ],
    bukti_intro=(
        "Bukti berupa hasil pipeline yang tampil sebagai dataset siap pakai pada platform: data halal hasil crawl "
        "dan data event hasil pengayaan geocoding."
    ),
    bukti=[
        ("img", "08-halal-restoran.png", "Gambar 1. Hasil pipeline crawl data halal — dataset restoran/zona KHAS tersertifikasi (BPJPH/LPPOM MUI) dengan sumber & koordinat tercatat."),
        ("img", "07-event-geocode.png", "Gambar 2. Hasil pipeline geocoding — 804 baris event dengan kolom alamat, lat/lon, dan sumber; tersedia unduh CSV/XLSX."),
        ("wis", "Cuplikan hasil pipeline data wisatawan mancanegara (agregat per negara asal):"),
    ],
    rtl_peran=(
        "Pada pengembangan platform, Data Engineer berperan mengotomasi lapisan Data Gathering — membaca struktur "
        "sumber, mengekstraksi data, menstandarkan, dan menyajikannya sebagai tabel siap analisis — sehingga proses "
        "yang kini semi-manual dapat berjalan berkelanjutan dengan kontrol validasi yang tetap dapat diaudit."
    ),
    rtl_bukti=[],
    penutup=(
        "Rangkaian pipeline crawl, geocoding, dan penyajian mandiri telah berjalan dan menghasilkan dataset "
        "terstandar. Tahap berikutnya adalah mengonsolidasikan pipeline ini ke dalam lapisan Data Gathering "
        "platform secara otomatis."
    ),
))

# 6. Data Analyst
MEMBERS.append(dict(
    filename="06-Data-Analyst",
    jabatan="Tenaga Ahli Data Analyst (Intermediate)",
    kualifikasi="S1 Sistem Informasi/Teknik Informatika/Ilmu Komputer/Statistik, pengalaman minimal 6 tahun",
    pendahuluan_peran=(
        "Sebagai Data Analyst, tenaga ahli bertanggung jawab menyiapkan dataset untuk pelaporan dan visualisasi, "
        "menyusun rekap indikator, memastikan kualitas dan keakuratan data yang disajikan, serta memvalidasi "
        "kebutuhan informasi bersama pengguna data."
    ),
    uraian_tugas=[
        "Menyiapkan dataset untuk pelaporan dan visualisasi sesuai kebutuhan Dinas;",
        "Menyusun laporan rutin (bulanan, triwulanan, tahunan) dan laporan ad-hoc yang relevan;",
        "Memastikan kualitas dan keakuratan data yang disajikan pada antarmuka pelaporan dan dashboard;",
        "Berkolaborasi dengan pengguna data untuk memvalidasi kebutuhan informasi.",
    ],
    pelaksanaan_intro=(
        "Pada bulan kedua, tenaga ahli menyiapkan dan menjamin mutu dataset baru (halal/ramah muslim, wisman, "
        "kuliner) serta dataset event yang diperkuat, kemudian merekapnya menjadi indikator yang tampil pada "
        "dashboard. Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Penyiapan dataset deliverable baru",
         "Menyusun dataset siap serah: 6 dataset wisata ramah muslim/halal (a.l. 134 restoran/zona KHAS, 49 hotel, 43 mall, 43 inovasi, 24 produk kreatif, 16 RPH), dataset wisman (per bulan/negara/pintu masuk), serta kuliner dan artis Top Chart — semua dapat diunduh CSV/XLSX."),
        ("Penguatan dataset event",
         "Menyiapkan dataset Jumlah Pengunjung Event 2026 (804 baris) versi diperkuat: setiap baris dilengkapi alamat, koordinat, dan sumber; 12 event prioritas diverifikasi ulang (7 mencapai keyakinan tinggi)."),
        ("Penjaminan kualitas & akurasi",
         "Memverifikasi nomor sertifikat halal dan kecocokan alamat/koordinat; menandai baris yang belum terverifikasi sebagai perlu-ditinjau agar tidak menyesatkan indikator."),
        ("Penyusunan rekap & validasi kebutuhan",
         "Menyusun rekap indikator Pengalaman Budaya (kuliner, seni & pertunjukan, wisatawan mancanegara) dan menyelaraskan definisi kolom dengan kebutuhan indikator GCI/GPCI bersama Business Analyst."),
    ],
    bukti_intro=(
        "Bukti berupa cuplikan dataset deliverable baru (halal), dataset event yang diperkuat, dan rekap wisman."
    ),
    bukti=[
        ("hal", "Cuplikan dataset restoran halal (deliverable baru — tersertifikasi BPJPH/LPPOM MUI):"),
        ("evt", "Cuplikan dataset event yang diperkuat (alamat & koordinat hasil geocode):"),
        ("wis", "Rekap wisatawan mancanegara per negara asal (Semester I 2026):"),
    ],
    rtl_peran=(
        "Pada pengembangan platform, Data Analyst berperan memanfaatkan Business Data Mart dan AI Analyst untuk "
        "mempercepat penyusunan laporan rutin & ad-hoc, dengan menjaga setiap angka tetap terhubung ke data, "
        "metrik, dan query yang dapat diaudit."
    ),
    rtl_bukti=[],
    penutup=(
        "Dataset baru dan dataset yang diperkuat telah disiapkan, diverifikasi, dan direkap menjadi indikator. "
        "Tahap berikutnya adalah membakukan laporan rutin di atas data mart platform."
    ),
))

# 7. Business Analyst
MEMBERS.append(dict(
    filename="07-Business-Analyst-1",
    jabatan="Tenaga Ahli Business Analyst 1 (Intermediate)",
    kualifikasi="S1 Sistem Informasi/Teknik Informatika/Ilmu Komputer/Manajemen/Statistik, pengalaman minimal 6 tahun",
    pendahuluan_peran=(
        "Sebagai Business Analyst, tenaga ahli bertanggung jawab menerjemahkan indikator strategis menjadi "
        "kebutuhan data konkret, memetakan sumber data, menyusun gap analysis, serta mendokumentasikan requirement "
        "dan use case — termasuk mengaitkan pekerjaan bulan ini dengan posisi daya saing kota global."
    ),
    uraian_tugas=[
        "Melakukan pemetaan kondisi tata kelola data pariwisata saat ini (as-is) melalui wawancara, observasi, dan analisis dokumen;",
        "Merumuskan kondisi tata kelola data yang dituju (to-be) bersama pemangku kepentingan;",
        "Menyusun gap analysis as-is vs to-be beserta rekomendasi pemenuhan;",
        "Menerjemahkan kebutuhan bisnis menjadi spesifikasi tata kelola data (standar data, metadata, kebijakan MDM, kebijakan akses);",
        "Mendokumentasikan requirement, alur proses, dan use case pengelolaan data pariwisata;",
        "Berkoordinasi dengan pemangku kepentingan internal dan eksternal Dinas;",
        "Mendukung penyusunan roadmap dan rencana implementasi standarisasi data.",
    ],
    pelaksanaan_intro=(
        "Pada bulan kedua, tenaga ahli menerjemahkan dimensi GCI Pengalaman Budaya menjadi kebutuhan data konkret "
        "(kuliner, seni & pertunjukan, wisatawan mancanegara), menambahkan aksis strategis wisata ramah muslim/halal, "
        "serta merumuskan kebutuhan kedaulatan data yang mendorong platform berjalan mandiri. Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Spesifikasi kebutuhan data indikator",
         "Menerjemahkan dimensi GCI Pengalaman Budaya menjadi daftar variabel & sumber data: kuliner (Michelin), seni & pertunjukan (venue), dan wisatawan mancanegara (per bulan/negara/pintu masuk)."),
        ("Perluasan aksis strategis (halal tourism)",
         "Merumuskan kebutuhan data wisata ramah muslim/halal sebagai daya tarik strategis Jakarta, memetakan sumber (BPJPH/LPPOM MUI, situs resmi, media) dan atribut wajib (nomor sertifikat, lokasi, sumber)."),
        ("Gap analysis & rekomendasi kedaulatan data",
         "Mengidentifikasi kesenjangan ketergantungan pada layanan pihak ketiga dan merekomendasikan operasionalisasi platform pada infrastruktur mandiri Dinas sebagai pemenuhan aspek kedaulatan & keberlanjutan data."),
        ("Dokumentasi requirement & use case",
         "Mendokumentasikan use case pemanfaatan (mis. ‘sebaran venue seni & pertunjukan per wilayah’, ‘negara asal wisatawan mancanegara dominan’, ‘destinasi ramah muslim tersertifikasi’) dan menautkannya ke KPI Disparekraf."),
        ("Koordinasi & roadmap",
         "Menyelaraskan definisi indikator bersama Data Analyst dan mendukung pemutakhiran roadmap standardisasi data bersama Project Manager."),
    ],
    bukti_intro=(
        "Bukti berupa perwujudan indikator Pengalaman Budaya pada dashboard, katalog data sebagai basis tata kelola, "
        "dan dataset aksis strategis wisata ramah muslim."
    ),
    bukti=[
        ("img", "02-gci.png", "Gambar 1. Terjemahan dimensi GCI Pengalaman Budaya menjadi empat sub-indikator pariwisata pada dashboard."),
        ("img", "06-sdi-katalog.png", "Gambar 2. Katalog data (182 dataset) sebagai basis tata kelola data (as-is → to-be)."),
        ("img", "08-halal-restoran.png", "Gambar 3. Perwujudan aksis strategis wisata ramah muslim — dataset tersertifikasi halal dengan sumber terverifikasi."),
    ],
    rtl_peran=(
        "Pada pengembangan platform, Business Analyst berperan menjaga keterhubungan platform dengan KPI utama "
        "Disparekraf (kontribusi ekonomi kreatif terhadap PDRB dan jumlah wisatawan mancanegara) serta posisi "
        "Jakarta pada indeks daya saing kota global, dan melanjutkan penyusunan roadmap standardisasi data."
    ),
    rtl_bukti=[],
    penutup=(
        "Spesifikasi kebutuhan data indikator Pengalaman Budaya, aksis wisata ramah muslim, dan rekomendasi "
        "kedaulatan data telah tersusun dan terwujud pada platform. Tahap berikutnya adalah memformalkan dokumen "
        "standardisasi data dan gap analysis sebagai dasar implementasi lanjutan."
    ),
))

if __name__ == "__main__":
    for m in MEMBERS:
        build(m)
    print("done:", len(MEMBERS), "reports")
