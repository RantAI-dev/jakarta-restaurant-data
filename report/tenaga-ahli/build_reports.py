# -*- coding: utf-8 -*-
"""
Generator Laporan Pelaksanaan Tugas Tenaga Ahli (7 anggota).
Kegiatan: Penyediaan dan Pengelolaan Data Statistik Pariwisata Jakarta — TA 2026.
Menghasilkan .docx (python-docx). Konversi ke .pdf dilakukan terpisah via LibreOffice.
"""
import csv, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(BASE, "assets")
REPO = os.path.abspath(os.path.join(BASE, "..", ".."))
INK = RGBColor(0x1a, 0x1a, 0x1a)
GREY = RGBColor(0x55, 0x55, 0x55)

# ---------- helpers ----------

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table, color="7f7f7f", sz="4"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)

def cell_text(cell, text, bold=False, size=10, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align: p.alignment = align
    run = p.add_run(text)
    run.bold = bold; run.font.size = Pt(size); run.font.name = "Arial"; run.font.color.rgb = color
    return p

def para(doc, text, size=11, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         color=INK, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after); pf.space_before = Pt(space_before); pf.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    run.font.name = "Arial"; run.font.color.rgb = color
    return p

def heading(doc, text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(size); run.font.name = "Arial"; run.font.color.rgb = INK
    return p

def bullets(doc, items, letters=False):
    for i, it in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.9); p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        prefix = f"{chr(97+i)}.  " if letters else "•  "
        r = p.add_run(prefix); r.font.name="Arial"; r.font.size=Pt(11); r.bold = letters
        r2 = p.add_run(it); r2.font.name="Arial"; r2.font.size=Pt(11); r2.font.color.rgb = INK

def add_image(doc, filename, caption, width_cm=15.5):
    path = os.path.join(ASSETS, filename)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Cm(width_cm))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9)
    r.font.name = "Arial"; r.font.color.rgb = GREY

def realisasi_table(doc, rows):
    """rows: list of (tugas, realisasi). Two-column mapping table."""
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_table_borders(t)
    hdr = t.rows[0].cells
    cell_text(hdr[0], "Uraian Tugas (KAK)", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_text(hdr[1], "Realisasi Bulan Ini", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    for c in hdr: set_cell_bg(c, "d9d9d9")
    for tugas, real in rows:
        r = t.add_row().cells
        cell_text(r[0], tugas, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
        cell_text(r[1], real, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    t.columns[0].width = Cm(6.6); t.columns[1].width = Cm(9.0)
    for row in t.rows:
        row.cells[0].width = Cm(6.6); row.cells[1].width = Cm(9.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def data_table(doc, header, rows, widths):
    t = doc.add_table(rows=1, cols=len(header)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; set_table_borders(t)
    for i, h in enumerate(header):
        cell_text(t.rows[0].cells[i], h, bold=True, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_bg(t.rows[0].cells[i], "d9d9d9")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cell_text(cells[i], v, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ---------- data samples from real TSV ----------

def tsv_rows(path, cols_idx, n=6):
    with open(path, encoding='utf-8') as f:
        r = list(csv.reader(f, delimiter='\t'))
    return [[r[k][i] for i in cols_idx] for k in range(1, 1+n)]

REST_TSV = os.path.join(REPO, "data-restoran-GCI-jakarta.tsv")
EVT_TSV = os.path.join(REPO, "data-pertunjukan-GCI-jakarta-2025.tsv")

REST_HEADER = ["No.", "Nama", "Jenis Cuisine", "Rating", "Ulasan", "Sumber"]
REST_ROWS = tsv_rows(REST_TSV, [0,1,2,4,5,6], 6)
REST_WIDTHS = [1.0, 4.2, 3.6, 1.6, 2.0, 2.1]

EVT_HEADER = ["No.", "Penyelenggara", "Nama Pertunjukan", "Tanggal", "Tempat", "Jenis"]
EVT_ROWS = tsv_rows(EVT_TSV, [0,1,2,3,4,5], 6)
EVT_WIDTHS = [0.9, 2.7, 3.9, 1.9, 3.0, 3.1]

# ---------- document skeleton ----------

def build(doc_meta):
    doc = Document()
    # base style
    st = doc.styles['Normal']; st.font.name = "Arial"; st.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.2); section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.2)

    # KOP
    para(doc, "PEMERINTAH PROVINSI DAERAH KHUSUS IBUKOTA JAKARTA", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, "DINAS PARIWISATA DAN EKONOMI KREATIF", size=13, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, "Bidang Data, Informasi dan Pengembangan Destinasi", size=10, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=2)
    # rule
    hr = doc.add_paragraph(); hr.paragraph_format.space_after = Pt(10)
    pPr = hr._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'12')
    bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),'1a1a1a'); pbdr.append(bottom); pPr.append(pbdr)

    para(doc, "LAPORAN PELAKSANAAN TUGAS TENAGA AHLI", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    para(doc, "Kegiatan Penyediaan dan Pengelolaan Data Statistik Pariwisata Jakarta — Tahun Anggaran 2026",
         size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=10)

    # IDENTITAS table
    t = doc.add_table(rows=0, cols=2); set_table_borders(t, color="bfbfbf")
    ident = [
        ("Nama Tenaga Ahli", "[Nama Tenaga Ahli]"),
        ("NIK / No. Kontrak", "[NIK / No. Kontrak]"),
        ("Jabatan / Posisi", doc_meta["jabatan"]),
        ("Kualifikasi", doc_meta["kualifikasi"]),
        ("Periode Laporan", "Bulan ke-1 (Laporan Bulanan Pertama)"),
        ("Sub Kegiatan", "Perencanaan Daya Tarik Wisata Provinsi"),
        ("Lokasi", "Provinsi DKI Jakarta"),
    ]
    for k, v in ident:
        row = t.add_row().cells
        cell_text(row[0], k, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT); set_cell_bg(row[0], "eeeeee")
        cell_text(row[1], v, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        row[0].width = Cm(4.8); row[1].width = Cm(10.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # I. PENDAHULUAN
    heading(doc, "I.  PENDAHULUAN")
    para(doc, PENDAHULUAN_UMUM)
    para(doc, doc_meta["pendahuluan_peran"])

    # II. URAIAN TUGAS
    heading(doc, "II.  URAIAN TUGAS")
    para(doc, "Sesuai Kerangka Acuan Kerja (KAK) kegiatan, uraian tugas untuk posisi ini adalah sebagai berikut:",
         space_after=4)
    bullets(doc, doc_meta["uraian_tugas"], letters=True)

    # III. PELAKSANAAN PEKERJAAN
    heading(doc, "III.  PELAKSANAAN PEKERJAAN BULAN INI")
    para(doc, doc_meta["pelaksanaan_intro"])
    realisasi_table(doc, doc_meta["realisasi"])

    # IV. BUKTI PELAKSANAAN
    heading(doc, "IV.  BUKTI PELAKSANAAN")
    para(doc, doc_meta["bukti_intro"], space_after=6)
    for item in doc_meta["bukti"]:
        if item[0] == "img":
            add_image(doc, item[1], item[2])
        elif item[0] == "rest":
            para(doc, item[1], bold=True, size=10.5, space_before=4, space_after=3,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
            data_table(doc, REST_HEADER, REST_ROWS, REST_WIDTHS)
            para(doc, "Sumber: berkas deliverable data-restoran-GCI-jakarta.tsv (84 baris kurasi).",
                 size=8.5, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
        elif item[0] == "evt":
            para(doc, item[1], bold=True, size=10.5, space_before=4, space_after=3,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
            data_table(doc, EVT_HEADER, EVT_ROWS, EVT_WIDTHS)
            para(doc, "Sumber: berkas deliverable data-pertunjukan-GCI-jakarta-2025.tsv (170 baris, 2025).",
                 size=8.5, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)

    # V. RENCANA TINDAK LANJUT
    heading(doc, "V.  RENCANA TINDAK LANJUT — PENGEMBANGAN PLATFORM AI-DATA")
    para(doc, RTL_UMUM)
    para(doc, doc_meta["rtl_peran"])
    for item in doc_meta.get("rtl_bukti", []):
        add_image(doc, item[0], item[1])

    # VI. PENUTUP
    heading(doc, "VI.  PENUTUP")
    para(doc, doc_meta["penutup"])

    # tanda tangan
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    sig = doc.add_table(rows=0, cols=2)
    r = sig.add_row().cells
    cell_text(r[0], "Mengetahui,\nPejabat Pembuat Komitmen", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(r[1], "Jakarta, ......................... 2026\nTenaga Ahli yang bersangkutan",
              size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4): sig.add_row()
    r = sig.add_row().cells
    cell_text(r[0], "(  Bima Agung  )", bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(r[1], "(  [Nama Tenaga Ahli]  )", bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = sig.add_row().cells
    cell_text(r[0], "NIP. 197907162011011008", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
    cell_text(r[1], doc_meta["jabatan"], size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)

    out = os.path.join(BASE, doc_meta["filename"] + ".docx")
    doc.save(out)
    print("wrote", out)

# ---------- shared narrative ----------

PENDAHULUAN_UMUM = (
    "Laporan ini disusun sebagai pertanggungjawaban pelaksanaan tugas tenaga ahli pada kegiatan "
    "Penyediaan dan Pengelolaan Data Statistik Pariwisata Jakarta Tahun Anggaran 2026, yang bertujuan "
    "mewujudkan tata kelola data pariwisata sesuai prinsip Satu Data Indonesia. Fokus pekerjaan pada "
    "bulan pertama ini adalah membangun fondasi data indikator Global City Index (GCI) — khususnya "
    "inventarisasi restoran & keberagaman kuliner internasional serta pertunjukan musik dan acara budaya "
    "di DKI Jakarta — beserta antarmuka dashboard awal untuk menampilkannya."
)

RTL_UMUM = (
    "Sebagai kelanjutan pekerjaan, tim mengusulkan pengembangan Jakarta Tourism Intelligence Platform, "
    "yaitu satu fondasi AI-Data yang menggabungkan lima kemampuan inti: Data Gathering (pengumpulan data "
    "dari web, database, file, dokumen, survei & sumber publik), Lakehouse Making (penyusunan data menjadi "
    "lapisan raw, clean, dan business data mart), Dashboarding (visualisasi indikator strategis), AI Analyst "
    "(bertanya langsung kepada data dengan bahasa natural, jawaban tetap dapat diaudit), dan Policy Insight "
    "Layer (fondasi analisis dampak, nowcasting & simulasi kebijakan). Platform ini menjadi wadah agar "
    "pekerjaan data yang saat ini masih dilakukan secara manual dapat berjalan terintegrasi, berkelanjutan, "
    "dan mendukung posisi Jakarta sebagai kota global."
)

# ---------- per-member content ----------

MEMBERS = []

# 1. Project Manager Analisis Data & Dashboard
MEMBERS.append(dict(
    filename="01-PM-Analisis-Data-Dashboard",
    jabatan="Project Manager Analisis Data dan Pembangunan Dashboard Pariwisata",
    kualifikasi="S2 / setara, pengalaman minimal 5 tahun",
    pendahuluan_peran=(
        "Sebagai Project Manager Analisis Data dan Pembangunan Dashboard Pariwisata, tenaga ahli bertanggung "
        "jawab merencanakan dan mengoordinasikan seluruh sumber daya tim, menetapkan lingkup indikator yang "
        "dikerjakan, memantau progres, mengelola risiko, serta memastikan output bulan pertama selesai sesuai "
        "standar kualitas dan harapan pemilik program."
    ),
    uraian_tugas=[
        "Mengelola dan merencanakan SDM yang dibutuhkan untuk kegiatan Analisis Data dan Pembangunan Dashboard Pariwisata, termasuk analisis, desain, pelaksanaan, pengujian/validasi, dan output yang dihasilkan;",
        "Memberikan laporan kegiatan Analisis Data dan Pembangunan Dashboard Pariwisata berupa progres pekerjaan dan pembaruan;",
        "Memastikan kegiatan selesai tepat waktu, sesuai anggaran, memenuhi standar kualitas, dan harapan pemilik program;",
        "Mengelola risiko dan isu pekerjaan, memastikan penyediaan informasi tepat waktu, serta melakukan mitigasi risiko dan langkah eskalasi.",
    ],
    pelaksanaan_intro=(
        "Pada bulan pertama, tenaga ahli memimpin penetapan lingkup pekerjaan pada dua indikator GCI prioritas "
        "(restoran & kuliner internasional; pertunjukan & acara budaya), membagi tugas antar anggota tim, dan "
        "mengonsolidasikan hasil menjadi tiga keluaran nyata: dua dataset deliverable, satu antarmuka dashboard "
        "publik, dan satu dokumen konsep platform. Rincian realisasi terhadap uraian tugas sebagai berikut:"
    ),
    realisasi=[
        ("Perencanaan SDM & lingkup pekerjaan",
         "Menetapkan lingkup dua indikator GCI dan membagi peran tim (BA, DBA, MDM, Data Engineer, Data Analyst, BI Developer); menyusun urutan kerja pengumpulan → pembersihan → penyajian data."),
        ("Pelaporan progres & pembaruan",
         "Menyusun laporan bulanan tim serta rekapitulasi capaian: 2.577 entri restoran (1.793 ber-rating) dan 308 entri pertunjukan/acara budaya termonitor pada dashboard."),
        ("Ketepatan waktu, mutu & anggaran",
         "Memastikan deliverable dataset restoran selesai pada tenggat serta mengendalikan biaya enumerasi data (pemanfaatan sumber gratis OSM dan pembatasan pemanggilan Places API agar tetap dalam kuota bebas biaya)."),
        ("Manajemen risiko & isu",
         "Mengidentifikasi risiko kualitas data (mis. rating salah-tautan pada mall/landmark) dan menetapkan mitigasi berupa verifikasi ulang serta pembersihan data sebelum dipublikasikan."),
    ],
    bukti_intro=(
        "Berikut bukti keluaran kegiatan yang berada di bawah koordinasi Project Manager, berupa tangkapan layar "
        "antarmuka publik dan rekapitulasi indikator."
    ),
    bukti=[
        ("img", "home-viewport.png", "Gambar 1. Halaman utama direktori data pariwisata DKI Jakarta (Jakarta Atlas)."),
        ("img", "gci-viewport.png", "Gambar 2. Dashboard Restoran GCI — rekap 2.577 entri, 1.793 ber-rating, 138 restoran hotel."),
        ("img", "events-viewport.png", "Gambar 3. Dashboard Pertunjukan & Budaya — 308 acara (125 konser, 24 festival, 159 budaya)."),
    ],
    rtl_peran=(
        "Dalam pengembangan platform, Project Manager berperan menyusun roadmap implementasi bertahap: Fase 1 "
        "(fondasi data & dashboard awal), Fase 2 (Data Platform Pilot, AI Analyst & semantic metrics), dan Fase 3 "
        "(integrasi data administratif, knowledge graph & governance); serta memastikan setiap fase selaras dengan "
        "KPI Disparekraf dan indikator daya saing kota global."
    ),
    rtl_bukti=[
        ("deck-09-roadmap.png", "Gambar 4. Roadmap implementasi bertahap platform (Fase 1–3)."),
    ],
    penutup=(
        "Secara keseluruhan, pekerjaan bulan pertama telah menghasilkan fondasi data dan dashboard awal sesuai "
        "rencana. Koordinasi tim akan dilanjutkan untuk memperluas cakupan indikator dan menyiapkan pengembangan "
        "platform pada tahap berikutnya."
    ),
))

# 2. Database Administrator
MEMBERS.append(dict(
    filename="02-Database-Administrator",
    jabatan="Tenaga Ahli Database Administrator (Senior)",
    kualifikasi="S1 Teknik Informatika/Ilmu Komputer/Sistem Informasi, pengalaman minimal 7 tahun",
    pendahuluan_peran=(
        "Sebagai Database Administrator, tenaga ahli bertanggung jawab merancang struktur penyimpanan data "
        "statistik pariwisata, menyusun skema tabel dan aturan integritas data, melaksanakan loading data hasil "
        "pengumpulan, serta mendokumentasikan struktur data agar konsisten dan dapat digunakan ulang."
    ),
    uraian_tugas=[
        "Merancang struktur basis data statistik pariwisata pada platform Database (logical & physical data model);",
        "Menyusun skema tabel, indeks, constraint, view, dan partisi sesuai kebutuhan analitik;",
        "Melaksanakan loading data hasil pengumpulan, pembersihan, dan validasi ke dalam Database;",
        "Melakukan optimasi kinerja basis data (query tuning, indexing, statistik, partitioning);",
        "Menyusun tata kelola akses dan keamanan basis data (role, privilege, audit trail);",
        "Menyusun pedoman backup, recovery, dan pemeliharaan basis data;",
        "Mendokumentasikan basis data (ERD, kamus tabel, panduan operasional).",
    ],
    pelaksanaan_intro=(
        "Pada bulan pertama, tenaga ahli merancang struktur data untuk dua domain (restoran dan pertunjukan) dalam "
        "bentuk skema tabel terstandar dengan kolom, tipe, dan sumber yang konsisten, serta menyiapkan lapisan data "
        "raw–clean–mart sebagai kerangka lakehouse. Sesuai lingkup KAK, pekerjaan difokuskan pada perancangan dan "
        "struktur data (bukan administrasi basis data operasional). Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Rancangan logical & physical model",
         "Menetapkan skema entitas restoran (Nama, Jenis Cuisine, Alamat, Rating, Jumlah Ulasan, Sumber, Tier/Area) dan pertunjukan (Penyelenggara, Nama, Tanggal, Tempat, Jenis, Jumlah Pengunjung, Keterangan, Sumber)."),
        ("Skema tabel, constraint & view",
         "Menyusun kolom terstandar dan aturan pengisian (mis. Jumlah Pengunjung hanya berisi angka, kosong bila tidak tersedia; penggunaan kapasitas sebagai proksi diberi keterangan)."),
        ("Loading data terkurasi & tervalidasi",
         "Memuat 84 baris restoran deliverable dan 308 entri pertunjukan ke struktur data aplikasi, termasuk kolom turunan alamat, harga, dan rating."),
        ("Optimasi & deduplikasi",
         "Menerapkan mekanisme deduplikasi lintas-sumber (normalisasi nama + tahun-bulan) sehingga entri ganda dan wilayah non-DKI tersaring sebelum dimuat."),
        ("Tata kelola akses & sumber",
         "Menandai sumber setiap nilai (Google/OSM/editorial) sebagai jejak keterlacakan (audit trail) pada setiap baris data."),
        ("Dokumentasi struktur data",
         "Menyusun kamus kolom dan pedoman pengisian yang menyertai kedua berkas deliverable."),
    ],
    bukti_intro=(
        "Bukti berupa struktur tabel data yang tampil pada dashboard dan cuplikan skema data deliverable."
    ),
    bukti=[
        ("img", "gci-viewport.png", "Gambar 1. Struktur tabel restoran pada dashboard — kolom terstandar Nama, Jenis Cuisine, Alamat, Harga, Rating, Jumlah Ulasan, Sumber, Tier."),
        ("rest", "Cuplikan skema & isi tabel restoran (deliverable):"),
        ("img", "deck-05-arsitektur.png", "Gambar 2. Rancangan arsitektur konseptual lapisan data: raw → clean → business data mart."),
    ],
    rtl_peran=(
        "Pada pengembangan platform, DBA berperan membangun lapisan Lakehouse (Raw Data Layer, Clean Data Layer, "
        "Business Data Mart) sebagai fondasi data yang terstandarisasi, tervalidasi, dan siap digunakan ulang oleh "
        "dashboard maupun AI Analyst, lengkap dengan arsip metadata dan jejak sumber."
    ),
    rtl_bukti=[],
    penutup=(
        "Struktur data untuk dua domain indikator telah dirancang, dimuat, dan didokumentasikan. Tahap berikutnya "
        "adalah memperluas skema untuk entitas pariwisata lain dan menerapkannya pada lapisan lakehouse platform."
    ),
))

# 3. MDM Specialist
MEMBERS.append(dict(
    filename="03-MDM-Specialist",
    jabatan="Tenaga Ahli Master Data Management (MDM) Specialist (Senior)",
    kualifikasi="S1 Sistem Informasi/Teknik Informatika/Ilmu Komputer/Statistik, pengalaman minimal 7 tahun",
    pendahuluan_peran=(
        "Sebagai MDM Specialist, tenaga ahli bertanggung jawab mengidentifikasi entitas master data pariwisata, "
        "menyusun strukturnya, serta melaksanakan konsolidasi, deduplikasi, verifikasi, dan validasi data dari "
        "berbagai sumber agar terbentuk data induk (golden record) yang tepercaya."
    ),
    uraian_tugas=[
        "Mengidentifikasi entitas master data pariwisata Jakarta dan menyusun struktur data master;",
        "Menyusun kebijakan dan prosedur MDM (data ownership, data stewardship, alur perubahan, persetujuan, dan audit);",
        "Melaksanakan pembentukan master data melalui konsolidasi, deduplikasi, verifikasi, dan validasi data dari berbagai sumber;",
        "Menyusun mekanisme pemeliharaan dan pemutakhiran master data secara berkala;",
        "Memastikan integrasi master data ke dalam database dan ketersediaannya bagi Dashboard melalui metadata.",
    ],
    pelaksanaan_intro=(
        "Pada bulan pertama, tenaga ahli mengidentifikasi entitas master (restoran, hotel/akomodasi, event/pertunjukan, "
        "dan kawasan) serta membentuk data induk melalui konsolidasi multi-sumber, deduplikasi, dan verifikasi. "
        "Pekerjaan penting yang menonjol adalah penyelarasan nama restoran di dalam hotel dan pembersihan rating yang "
        "salah-tautan. Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Identifikasi entitas & struktur master",
         "Menetapkan entitas master restoran, hotel/akomodasi, event, dan kawasan; menyusun atribut kunci dan relasi antar-entitas (mis. restoran ↔ hotel induk)."),
        ("Kebijakan & prosedur stewardship",
         "Menetapkan aturan kepemilikan nilai per-sumber dan prosedur perubahan (nilai hasil riset tidak pernah difabrikasi; bila tak terverifikasi, dikembalikan ke status kosong/‘belum ada rating’)."),
        ("Konsolidasi, deduplikasi & verifikasi",
         "Menggabungkan seed kurasi dengan data bulk OSM; menyelaraskan 112 dari 138 restoran-hotel agar menampilkan nama restoran, bukan sekadar nama hotel; memverifikasi kecocokan placeId."),
        ("Validasi & pembersihan golden record",
         "Membersihkan rating yang keliru menautkan restoran ke mall/landmark (mis. Plaza Senayan, Museum Sejarah Jakarta) sehingga tidak menyesatkan indikator."),
        ("Integrasi master ke dashboard",
         "Memastikan atribut master (nama, alamat, tier, sumber, harga) tersedia dan konsisten pada antarmuka dashboard."),
    ],
    bukti_intro=(
        "Bukti berupa tampilan entitas master pada dashboard beserta atribut sumber, dan cuplikan data induk restoran."
    ),
    bukti=[
        ("img", "gci-viewport.png", "Gambar 1. Entitas master restoran & restoran-hotel dengan atribut Tier dan Sumber; nama restoran diselaraskan dari nama hotel induk."),
        ("rest", "Cuplikan data induk (golden record) restoran:"),
    ],
    rtl_peran=(
        "Pada pengembangan platform, MDM Specialist berperan memastikan setiap sumber data yang masuk melalui Data "
        "Gathering dikonsolidasikan menjadi master data tepercaya sebelum dipublikasikan — sehingga dashboard dan AI "
        "Analyst bekerja di atas definisi entitas yang seragam, terverifikasi, dan dapat diaudit."
    ),
    rtl_bukti=[
        ("deck-04-platform5.png", "Gambar 2. Lima kemampuan inti platform; master data menjadi penjamin mutu antara Data Gathering dan Lakehouse."),
    ],
    penutup=(
        "Entitas master untuk domain restoran dan pertunjukan telah dibentuk dan dibersihkan. Tahap berikutnya adalah "
        "memformalkan kebijakan stewardship dan memperluas master data ke seluruh entitas pariwisata pada platform."
    ),
))

# 4. BI Developer
MEMBERS.append(dict(
    filename="04-BI-Developer-1",
    jabatan="Tenaga Ahli BI Developer 1 (Junior)",
    kualifikasi="S1 Sistem Informasi/Teknik Informatika/Ilmu Komputer/Statistik, pengalaman minimal 5 tahun",
    pendahuluan_peran=(
        "Sebagai BI Developer, tenaga ahli bertanggung jawab membangun dashboard interaktif yang menyajikan "
        "indikator pariwisata, menyusun visualisasi yang konsisten dengan definisi data, serta menyediakan fitur "
        "penyaringan dan ekspor bagi pengguna di lingkungan Dinas."
    ),
    uraian_tugas=[
        "Mengkonfigurasi koneksi Dashboard ke database melalui semantic layer/metadata MDM;",
        "Membangun dashboard interaktif (overview pariwisata, kunjungan wisatawan, akomodasi, destinasi, kontribusi ekonomi, dan dashboard tematik lainnya);",
        "Menyusun visualisasi yang konsisten dengan standar metadata dan definisi data;",
        "Mengatur akses pengguna, jadwal refresh data, dan pemeliharaan dashboard;",
        "Menyusun panduan pemanfaatan dashboard bagi pengguna di lingkungan Dinas.",
    ],
    pelaksanaan_intro=(
        "Pada bulan pertama, tenaga ahli membangun antarmuka dashboard publik yang menampilkan dua indikator GCI "
        "beserta ringkasan statistik, filter, dan ekspor. Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Koneksi dashboard ke sumber data",
         "Menghubungkan dashboard ke data store terstruktur (restoran & pertunjukan) melalui lapisan data aplikasi."),
        ("Dashboard interaktif tematik",
         "Membangun dashboard Restoran GCI dan Pertunjukan & Budaya, lengkap dengan kartu ringkasan (total, ber-rating, hotel, cafe / konser, festival, budaya)."),
        ("Visualisasi konsisten dengan definisi",
         "Menambahkan kolom Harga (tingkat harga, penanda * untuk nilai editorial), Rating, dan Jumlah Ulasan sesuai definisi metadata; penanda verifikasi untuk nilai perkiraan/kapasitas."),
        ("Filter, pencarian & ekspor",
         "Menyediakan filter kategori (Hotel ★4/★3, Restoran, Cafe; konser/festival/budaya), filter tahun (2025/2026), pencarian, dan ekspor ke Excel."),
        ("Panduan pemanfaatan",
         "Menyusun antarmuka dwibahasa (ID/EN) agar dashboard mudah dimanfaatkan pengguna Dinas."),
    ],
    bukti_intro=(
        "Bukti berupa tangkapan layar dashboard interaktif yang dibangun."
    ),
    bukti=[
        ("img", "gci-viewport.png", "Gambar 1. Dashboard Restoran GCI — kartu ringkasan, filter kategori/rating, kolom Harga & Rating, dan tombol Ekspor Excel."),
        ("img", "events-viewport.png", "Gambar 2. Dashboard Pertunjukan & Budaya — filter tahun & kategori, ringkasan konser/festival/budaya."),
        ("img", "home-viewport.png", "Gambar 3. Halaman utama direktori sebagai pintu masuk seluruh dashboard tematik."),
    ],
    rtl_peran=(
        "Pada pengembangan platform, BI Developer berperan memperluas lapisan Dashboarding menjadi indikator strategis "
        "penuh (performa destinasi, event & MICE, tren sentimen, okupansi hotel, aktivitas kawasan, efektivitas promosi) "
        "serta menyiapkan integrasi dengan AI Analyst agar pengguna dapat bertanya langsung kepada data."
    ),
    rtl_bukti=[
        ("deck-04-platform5.png", "Gambar 4. Posisi Dashboarding di antara lima kemampuan inti platform."),
    ],
    penutup=(
        "Dua dashboard indikator GCI telah dibangun dan dapat diakses publik. Tahap berikutnya adalah menambah "
        "dashboard tematik dan menautkannya dengan lapisan AI Analyst pada platform."
    ),
))

# 5. Data Engineer
MEMBERS.append(dict(
    filename="05-Data-Engineer-2",
    jabatan="Tenaga Ahli Data Engineer 2 (Junior)",
    kualifikasi="S1 Teknik Informatika/Ilmu Komputer/Sistem Informasi, pengalaman minimal 6 tahun",
    pendahuluan_peran=(
        "Sebagai Data Engineer, tenaga ahli bertanggung jawab merancang dan menjalankan pipeline pengumpulan dan "
        "integrasi data (ETL/ELT) dari berbagai sumber, serta membersihkan, mentransformasi, dan menstandarisasi "
        "raw data agar siap dianalisis dan divisualisasikan."
    ),
    uraian_tugas=[
        "Merancang dan melaksanakan pipeline pengumpulan dan integrasi data (ETL/ELT) dari berbagai sumber ke dalam Database;",
        "Membersihkan, mentransformasi, dan menstandarisasi raw data agar siap digunakan untuk analisis dan visualisasi;",
        "Memastikan integrasi data multi-sumber dilakukan dengan menjaga kualitas dan keterlacakan data;",
        "Menerapkan standar metadata dan master data dalam proses integrasi data;",
        "Menjamin keamanan data dan kepatuhan terhadap tata kelola data pada proses integrasi.",
    ],
    pelaksanaan_intro=(
        "Pada bulan pertama, tenaga ahli membangun rangkaian pipeline pengumpulan data dari sumber terbuka dan "
        "layanan peta, lalu membersihkan dan menstandarkannya menjadi dataset siap pakai. Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Pipeline pengumpulan & integrasi",
         "Membangun pipeline penarikan data restoran dari OpenStreetMap (Overpass) yang dipotong tepat pada batas administrasi DKI Jakarta, serta pipeline penyusun data pertunjukan 2026 dari riset web paralel."),
        ("Pembersihan & standardisasi",
         "Menormalkan tanggal ke format ISO, angka pengunjung ke format standar, serta menstandarkan alamat dan penamaan; menyaring wilayah non-DKI (Bodetabek)."),
        ("Integrasi multi-sumber & keterlacakan",
         "Menggabungkan data OSM, Google Places (rating/alamat), dan riset editorial dengan menandai sumber tiap nilai untuk menjaga keterlacakan."),
        ("Penerapan standar metadata & master",
         "Menerapkan struktur kolom dan aturan master data (mis. override nama restoran-hotel, tingkat harga) secara konsisten dalam proses integrasi."),
        ("Kepatuhan & efisiensi biaya",
         "Mengutamakan sumber gratis dan membatasi pemanggilan API berbayar agar tetap dalam kuota bebas biaya, sekaligus menjaga kepatuhan penggunaan data."),
    ],
    bukti_intro=(
        "Bukti berupa hasil integrasi data yang tampil pada dashboard dan cuplikan dataset pertunjukan hasil pipeline."
    ),
    bukti=[
        ("img", "events-viewport.png", "Gambar 1. Hasil pipeline pengumpulan pertunjukan — 308 entri terintegrasi (2025–2026) dengan sumber tercatat."),
        ("evt", "Cuplikan dataset pertunjukan hasil pipeline (dengan kolom Sumber untuk keterlacakan):"),
        ("img", "deck-05-arsitektur.png", "Gambar 2. Alur konseptual dari sumber data → ingestion (AI Data Gathering) → lakehouse → output."),
    ],
    rtl_peran=(
        "Pada pengembangan platform, Data Engineer berperan mengotomasi lapisan Data Gathering — membaca struktur "
        "sumber, mengenali pola, mengekstraksi data, dan menyarankan bentuk tabel siap dianalisis — sehingga proses "
        "yang kini manual dapat berjalan berkelanjutan dengan tetap dapat dikontrol manual (validasi, struktur tabel, SQL)."
    ),
    rtl_bukti=[],
    penutup=(
        "Rangkaian pipeline pengumpulan dan integrasi data untuk dua domain telah berjalan dan menghasilkan dataset "
        "terstandar. Tahap berikutnya adalah mengonsolidasikan pipeline ini ke dalam lapisan Data Gathering platform."
    ),
))

# 6. Data Analyst
MEMBERS.append(dict(
    filename="06-Data-Analyst",
    jabatan="Tenaga Ahli Data Analyst (Intermediate)",
    kualifikasi="S1 Sistem Informasi/Teknik Informatika/Ilmu Komputer/Statistik, pengalaman minimal 6 tahun",
    pendahuluan_peran=(
        "Sebagai Data Analyst, tenaga ahli bertanggung jawab menyiapkan dataset untuk pelaporan dan visualisasi, "
        "menyusun laporan, memastikan kualitas dan keakuratan data yang disajikan, serta berkolaborasi dengan "
        "pengguna data untuk memvalidasi kebutuhan informasi."
    ),
    uraian_tugas=[
        "Menyiapkan dataset untuk pelaporan dan visualisasi sesuai kebutuhan Dinas;",
        "Menyusun laporan rutin (bulanan, triwulanan, tahunan) dan laporan ad-hoc yang relevan;",
        "Memastikan kualitas dan keakuratan data yang disajikan pada antarmuka pelaporan dan dashboard;",
        "Berkolaborasi dengan pengguna data untuk memvalidasi kebutuhan informasi.",
    ],
    pelaksanaan_intro=(
        "Pada bulan pertama, tenaga ahli menyiapkan dua dataset deliverable yang rapi dan terverifikasi, serta "
        "melakukan penjaminan mutu atas data yang ditampilkan pada dashboard. Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Penyiapan dataset deliverable",
         "Menyusun berkas deliverable restoran (84 baris terkurasi, lengkap kolom Alamat) dan pertunjukan (170 baris 2025 + 144 baris 2026) dalam format tabel siap serah dan siap ekspor Excel."),
        ("Penyusunan laporan & rekap",
         "Menyusun rekap indikator: 2.577 entri restoran (1.793 ber-rating; 138 restoran hotel; 712 cafe) dan 308 acara (125 konser, 24 festival, 159 budaya)."),
        ("Penjaminan kualitas & akurasi",
         "Memverifikasi kecocokan rating terhadap venue yang benar dan menandai nilai perkiraan/kapasitas dengan penanda khusus agar tidak menyesatkan."),
        ("Validasi kebutuhan informasi",
         "Menyelaraskan kolom dan definisi data dengan kebutuhan indikator GCI/GPCI bersama Business Analyst dan pengguna Dinas."),
    ],
    bukti_intro=(
        "Bukti berupa cuplikan dua dataset deliverable dan rekap indikator pada dashboard."
    ),
    bukti=[
        ("rest", "Cuplikan dataset restoran (deliverable, terurut berdasarkan rating & jumlah ulasan):"),
        ("evt", "Cuplikan dataset pertunjukan (deliverable 2025):"),
        ("img", "events-viewport.png", "Gambar 1. Rekap indikator pertunjukan pada dashboard, konsisten dengan dataset deliverable."),
    ],
    rtl_peran=(
        "Pada pengembangan platform, Data Analyst berperan memanfaatkan Business Data Mart dan AI Analyst untuk "
        "mempercepat penyusunan laporan rutin dan ad-hoc, sekaligus menjaga agar setiap angka yang dilaporkan tetap "
        "terhubung ke data, metrik, dan query yang dapat diaudit."
    ),
    rtl_bukti=[],
    penutup=(
        "Dua dataset deliverable telah disiapkan, diverifikasi, dan direkap menjadi indikator. Tahap berikutnya adalah "
        "membakukan laporan rutin di atas data mart platform."
    ),
))

# 7. Business Analyst
MEMBERS.append(dict(
    filename="07-Business-Analyst-1",
    jabatan="Tenaga Ahli Business Analyst 1 (Intermediate)",
    kualifikasi="S1 Sistem Informasi/Teknik Informatika/Ilmu Komputer/Manajemen/Statistik, pengalaman minimal 6 tahun",
    pendahuluan_peran=(
        "Sebagai Business Analyst, tenaga ahli bertanggung jawab memetakan kondisi tata kelola data saat ini (as-is), "
        "merumuskan kondisi yang dituju (to-be), menyusun gap analysis, menerjemahkan kebutuhan bisnis menjadi "
        "spesifikasi data, serta mendokumentasikan requirement dan use case."
    ),
    uraian_tugas=[
        "Melakukan pemetaan kondisi tata kelola data pariwisata saat ini (as-is) melalui wawancara, observasi, dan analisis dokumen;",
        "Merumuskan kondisi tata kelola data yang dituju (to-be) bersama pemangku kepentingan;",
        "Menyusun gap analysis as-is vs to-be beserta rekomendasi pemenuhan;",
        "Menerjemahkan kebutuhan bisnis menjadi spesifikasi tata kelola data (standar data, metadata, kebijakan MDM, kebijakan akses);",
        "Mendokumentasikan requirement, alur proses, dan use case pengelolaan data pariwisata;",
        "Berkoordinasi dengan pemangku kepentingan internal dan eksternal Dinas;",
        "Mendukung penyusunan roadmap dan rencana implementasi standarisasi data.",
    ],
    pelaksanaan_intro=(
        "Pada bulan pertama, tenaga ahli memetakan kondisi data pariwisata yang masih terfragmentasi (as-is), "
        "merumuskan kondisi decision intelligence yang dituju (to-be), dan menerjemahkan indikator GCI/GPCI menjadi "
        "kebutuhan data konkret. Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Pemetaan kondisi as-is",
         "Mengidentifikasi masalah utama: data tersebar di banyak sumber, laporan historis dan terlambat, analisis manual, dan belum ada sistem yang menghubungkan data hingga aksi."),
        ("Perumusan kondisi to-be",
         "Merumuskan visi alur ‘data terfragmentasi → terintegrasi → insight → prediksi → simulasi → aksi kebijakan’ yang proaktif, terukur, dan dapat diaudit."),
        ("Gap analysis & rekomendasi",
         "Menyusun rekomendasi fondasi intelligence end-to-end (bukan sekadar dashboard) sebagai pemenuhan kesenjangan."),
        ("Spesifikasi kebutuhan data indikator",
         "Menerjemahkan indikator GPCI Cultural Interaction (konferensi internasional, event budaya, museum, teater, kuliner, kamar hotel, dsb.) menjadi daftar variabel dan sumber data yang harus dikumpulkan."),
        ("Dokumentasi requirement & use case",
         "Mendokumentasikan use case AI Analyst (mis. ‘kawasan mana yang menunjukkan peningkatan minat wisatawan?’, ‘event mana yang paling berdampak terhadap okupansi hotel?’) dan menautkannya ke KPI Disparekraf."),
    ],
    bukti_intro=(
        "Bukti berupa pemetaan indikator GPCI ke sumber data dan rancangan alur solusi, serta perwujudan awal pada dashboard."
    ),
    bukti=[
        ("img", "deck-12-gpci.png", "Gambar 1. Pemetaan indikator pendukung GPCI ke sumber data (as-is → kebutuhan data)."),
        ("img", "deck-04-platform5.png", "Gambar 2. Terjemahan kebutuhan bisnis menjadi lima kemampuan inti platform (to-be)."),
        ("img", "gci-viewport.png", "Gambar 3. Perwujudan awal kebutuhan data indikator restoran GCI pada dashboard."),
    ],
    rtl_peran=(
        "Pada pengembangan platform, Business Analyst berperan menjaga keterhubungan platform dengan dua KPI utama "
        "Disparekraf (kontribusi ekonomi kreatif terhadap PDRB dan jumlah wisatawan mancanegara) serta posisi Jakarta "
        "pada indeks daya saing kota global, dan menyusun roadmap standardisasi data bersama Project Manager."
    ),
    rtl_bukti=[
        ("deck-09-roadmap.png", "Gambar 4. Roadmap implementasi standardisasi & platform (Fase 1–3)."),
    ],
    penutup=(
        "Pemetaan as-is/to-be dan spesifikasi kebutuhan data indikator telah tersusun. Tahap berikutnya adalah "
        "memformalkan dokumen standardisasi data dan gap analysis sebagai dasar implementasi platform."
    ),
))

if __name__ == "__main__":
    for m in MEMBERS:
        build(m)
    print("done:", len(MEMBERS), "reports")
