# -*- coding: utf-8 -*-
"""
Generator Laporan Pelaksanaan Tugas Tenaga Ahli (7 anggota) — BULAN KE-3.
Kegiatan: Penyediaan dan Pengelolaan Data Statistik Pariwisata Jakarta — TA 2026.

Fokus bulan ke-3 (Agustus 2026):
  1) Lake House data pariwisata (Bronze–Silver–Gold) BERJALAN OTOMATIS HARIAN
     di server Dinas — menutup arahan MoM 13 Juli (lake house + fleksibel ke Oracle).
  2) Penerapan platform lakehouse Rantai Lake (produk platform yang sudah ada —
     BUKAN dikembangkan dalam kegiatan ini): disambungkan ke lake house Dinas
     sebagai alat pengelolaan + dashboarding + asisten data.
  3) Aplikasi data Dispar kini membaca 100% dari lakehouse (tidak lagi menembak
     API luar saat halaman dibuka) + pembuat grafik mandiri per dataset +
     dashboard Wisman Target vs Realisasi RPJMD 2025–2030.
  4) Dokumentasi REST API publik, buku "Statistika Pariwisata Perkotaan",
     dan penelusuran data kunjungan 31 DTW.

Angka pada laporan diverifikasi langsung dari lakehouse (ClickHouse) 26 Agu 2026.
Menghasilkan .docx (python-docx). Konversi ke .pdf via LibreOffice (langkah terpisah).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

BASE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(BASE, "assets-bulan3")
INK = RGBColor(0x1a, 0x1a, 0x1a)
GREY = RGBColor(0x55, 0x55, 0x55)

# ---------- helpers ----------

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table, color="7f7f7f", sz="4"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)

def cell_text(cell, text, bold=False, size=10, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align: p.alignment = align
    run = p.add_run(text)
    run.bold = bold; run.font.size = Pt(size); run.font.name = "Arial"; run.font.color.rgb = color
    return p

def para(doc, text, size=11, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         color=INK, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after); pf.space_before = Pt(space_before); pf.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    run.font.name = "Arial"; run.font.color.rgb = color
    return p

def heading(doc, text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(size); run.font.name = "Arial"; run.font.color.rgb = INK
    return p

def bullets(doc, items, letters=False):
    for i, it in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.9); p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        prefix = f"{chr(97+i)}.  " if letters else "•  "
        r = p.add_run(prefix); r.font.name = "Arial"; r.font.size = Pt(11); r.bold = letters
        r2 = p.add_run(it); r2.font.name = "Arial"; r2.font.size = Pt(11); r2.font.color.rgb = INK

def add_image(doc, filename, caption, width_cm=15.5, max_h_cm=19.0):
    path = os.path.join(ASSETS, filename)
    w_px, h_px = Image.open(path).size
    w = width_cm
    h = width_cm * h_px / w_px
    if h > max_h_cm:
        h = max_h_cm; w = max_h_cm * w_px / h_px
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Cm(w), height=Cm(h))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9)
    r.font.name = "Arial"; r.font.color.rgb = GREY

def realisasi_table(doc, rows):
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_table_borders(t)
    hdr = t.rows[0].cells
    cell_text(hdr[0], "Uraian Tugas (KAK)", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_text(hdr[1], "Realisasi Bulan Ini", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    for c in hdr: set_cell_bg(c, "d9d9d9")
    for tugas, real in rows:
        r = t.add_row().cells
        cell_text(r[0], tugas, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
        cell_text(r[1], real, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    t.columns[0].width = Cm(6.6); t.columns[1].width = Cm(9.0)
    for row in t.rows:
        row.cells[0].width = Cm(6.6); row.cells[1].width = Cm(9.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def data_table(doc, header, rows, widths):
    t = doc.add_table(rows=1, cols=len(header)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; set_table_borders(t)
    for i, h in enumerate(header):
        cell_text(t.rows[0].cells[i], h, bold=True, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_bg(t.rows[0].cells[i], "d9d9d9")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cell_text(cells[i], v, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ---------- data terverifikasi dari lakehouse (26 Agustus 2026) ----------

LAKE_HEADER = ["Lapisan", "Bentuk penyimpanan", "Isi", "Jumlah"]
LAKE_ROWS = [
    ["Bronze (mentah)", "Apache Iceberg di object storage", "Salinan apa adanya dataset SDI + berkas pendataan lapangan, lengkap kolom audit (waktu tarik, sumber, batch)", "225 tabel"],
    ["Silver (bersih)", "View bertipe di ClickHouse", "Kolom sudah bertipe (angka/tanggal/teks bersih) + dimensi bersama (negara, bulan)", "224 view"],
    ["Gold (penyaji)", "Tabel agregat ClickHouse", "Mart siap dashboard: wisman, kuliner, event, atlas, kesiapan GCI, kunjungan DTW", "6 mart"],
]
LAKE_WIDTHS = [3.0, 4.2, 6.4, 2.0]

MART_HEADER = ["Mart Gold", "Isi", "Baris"]
MART_ROWS = [
    ["serving.mart_wisman", "Wisatawan mancanegara per bulan / negara / kawasan", "276"],
    ["serving.mart_kunjungan_dtw", "Kunjungan 31+ daya tarik wisata (DTW) beserta sumber & tanggal terbit", "35"],
    ["serving.mart_gci_readiness", "Kesiapan data per indikator GCI/GPCI", "28"],
    ["serving.mart_event", "Event & pengunjung event", "9"],
    ["serving.mart_kuliner", "Kuliner (Michelin/TripAdvisor) per wilayah", "5"],
    ["serving.mart_atlas", "Pendataan lapangan Atlas (restoran, pertunjukan, suvenir, nightlife)", "4"],
]
MART_WIDTHS = [5.0, 8.2, 2.4]

QUAL_HEADER = ["Pemeriksaan (gerbang mutu harian)", "Hasil 26 Agu 2026"]
QUAL_ROWS = [
    ["Lolos (pass)", "223 pemeriksaan"],
    ["Peringatan (warn) — nilai gagal konversi di bawah ambang", "42 pemeriksaan"],
    ["Gagal (fail) — kolom melewati ambang, barisnya dikarantina", "83 pemeriksaan"],
    ["Baris masuk karantina (tidak ikut ke lapisan penyaji)", "124 kelompok temuan"],
    ["Riwayat jumlah baris tersimpan (deteksi data anjlok)", "3.130 catatan"],
]
QUAL_WIDTHS = [9.4, 6.2]

KARANTINA_HEADER = ["Alasan karantina", "Kolom terdampak", "Nilai gagal konversi"]
KARANTINA_ROWS = [
    ["Nilai bukan angka pada kolom jumlah", "9 tabel", "52.735"],
    ["Nilai bukan angka pada kolom lon (bujur)", "7 tabel", "195"],
    ["Nilai bukan angka pada kolom lat (lintang)", "7 tabel", "195"],
    ["Nilai bukan angka pada kolom jumlah_ulasan", "4 tabel", "86"],
    ["Nilai bukan angka pada kolom rating", "3 tabel", "81"],
]
KARANTINA_WIDTHS = [7.4, 4.0, 4.2]

WIS_HEADER = ["Kawasan asal", "Total kunjungan (mart Gold)"]
WIS_ROWS = [
    ["Asia", "1.310.934"],
    ["Lainnya", "363.751"],
    ["Eropa", "235.750"],
    ["Timur Tengah", "152.938"],
    ["Amerika", "85.815"],
    ["Oseania", "69.579"],
    ["Afrika", "4.685"],
]
WIS_WIDTHS = [7.8, 7.8]

DTW_HEADER = ["Destinasi", "Periode angka yang terbit", "Total", "Sumber"]
DTW_ROWS = [
    ["Kawasan Kota Tua Jakarta", "27 Jun – 5 Jul 2026 (lintas bulan)", "99.663", "AyoJakarta.com mengutip data UPK Kota Tua"],
    ["Taman Lapangan Banteng", "19 Juli 2026 (satu hari)", "14.000", "Siaran pers Pemprov DKI Jakarta (jakarta.go.id)"],
    ["29 destinasi lainnya", "Tidak ditemukan angka bulan penuh", "—", "Tidak ada publikasi bulanan dari pengelola"],
]
DTW_WIDTHS = [4.2, 4.4, 1.9, 5.1]

# ---------- document skeleton ----------

def build(doc_meta):
    doc = Document()
    st = doc.styles['Normal']; st.font.name = "Arial"; st.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.2); section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.2)

    # KOP
    para(doc, "PEMERINTAH PROVINSI DAERAH KHUSUS IBUKOTA JAKARTA", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, "DINAS PARIWISATA DAN EKONOMI KREATIF", size=13, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, "Bidang Data, Informasi dan Pengembangan Destinasi", size=10, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=2)
    hr = doc.add_paragraph(); hr.paragraph_format.space_after = Pt(10)
    pPr = hr._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '1a1a1a'); pbdr.append(bottom); pPr.append(pbdr)

    para(doc, "LAPORAN PELAKSANAAN TUGAS TENAGA AHLI", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    para(doc, "Kegiatan Penyediaan dan Pengelolaan Data Statistik Pariwisata Jakarta — Tahun Anggaran 2026",
         size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=10)

    # IDENTITAS
    t = doc.add_table(rows=0, cols=2); set_table_borders(t, color="bfbfbf")
    ident = [
        ("Nama Tenaga Ahli", "[Nama Tenaga Ahli]"),
        ("NIK / No. Kontrak", "[NIK / No. Kontrak]"),
        ("Jabatan / Posisi", doc_meta["jabatan"]),
        ("Kualifikasi", doc_meta["kualifikasi"]),
        ("Periode Laporan", "Bulan ke-3 (Laporan Bulanan Ketiga)"),
        ("Sub Kegiatan", "Perencanaan Daya Tarik Wisata Provinsi"),
        ("Lokasi", "Provinsi DKI Jakarta"),
    ]
    for k, v in ident:
        row = t.add_row().cells
        cell_text(row[0], k, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT); set_cell_bg(row[0], "eeeeee")
        cell_text(row[1], v, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        row[0].width = Cm(4.8); row[1].width = Cm(10.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    heading(doc, "I.  PENDAHULUAN")
    para(doc, PENDAHULUAN_UMUM)
    para(doc, doc_meta["pendahuluan_peran"])

    heading(doc, "II.  URAIAN TUGAS")
    para(doc, "Sesuai Kerangka Acuan Kerja (KAK) kegiatan, uraian tugas untuk posisi ini adalah sebagai berikut:",
         space_after=4)
    bullets(doc, doc_meta["uraian_tugas"], letters=True)

    heading(doc, "III.  PELAKSANAAN PEKERJAAN BULAN INI")
    para(doc, doc_meta["pelaksanaan_intro"])
    realisasi_table(doc, doc_meta["realisasi"])

    heading(doc, "IV.  BUKTI PELAKSANAAN")
    para(doc, doc_meta["bukti_intro"], space_after=6)
    for item in doc_meta["bukti"]:
        kind = item[0]
        if kind == "img":
            add_image(doc, item[1], item[2])
        else:
            spec = TABLES[kind]
            para(doc, item[1], bold=True, size=10.5, space_before=4, space_after=3,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
            data_table(doc, spec[0], spec[1], spec[2])
            para(doc, spec[3], size=8.5, italic=True, color=GREY,
                 align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)

    heading(doc, "V.  RENCANA TINDAK LANJUT — PENGEMBANGAN PLATFORM AI-DATA")
    para(doc, RTL_UMUM)
    para(doc, doc_meta["rtl_peran"])
    for item in doc_meta.get("rtl_bukti", []):
        add_image(doc, item[0], item[1])

    heading(doc, "VI.  PENUTUP")
    para(doc, doc_meta["penutup"])

    # tanda tangan
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    sig = doc.add_table(rows=0, cols=2)
    r = sig.add_row().cells
    cell_text(r[0], "Mengetahui,\nPejabat Pembuat Komitmen", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(r[1], "Jakarta, ......................... 2026\nTenaga Ahli yang bersangkutan",
              size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4): sig.add_row()
    r = sig.add_row().cells
    cell_text(r[0], "(  Bima Agung  )", bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(r[1], "(  [Nama Tenaga Ahli]  )", bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = sig.add_row().cells
    cell_text(r[0], "NIP. 197907162011011008", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
    cell_text(r[1], doc_meta["jabatan"], size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)

    outdir = os.path.join(BASE, "bulan-3")
    os.makedirs(outdir, exist_ok=True)
    out = os.path.join(outdir, doc_meta["filename"] + ".docx")
    doc.save(out)
    print("wrote", out)

TABLES = {
    "lake": (LAKE_HEADER, LAKE_ROWS, LAKE_WIDTHS,
             "Sumber: lakehouse Dinas per 26 Agustus 2026. Bronze menyimpan data mentah dalam format terbuka (Apache Iceberg) sehingga tidak terkunci pada satu mesin basis data."),
    "mart": (MART_HEADER, MART_ROWS, MART_WIDTHS,
             "Sumber: lapisan Gold lakehouse (ClickHouse) per 26 Agustus 2026 — seluruh angka pada dashboard dibaca dari tabel-tabel ini."),
    "qual": (QUAL_HEADER, QUAL_ROWS, QUAL_WIDTHS,
             "Sumber: gerbang mutu (quality gate) yang berjalan otomatis setiap hari antara lapisan Silver dan Gold; hasil pemeriksaan tersimpan dan dapat ditelusuri per tanggal."),
    "karantina": (KARANTINA_HEADER, KARANTINA_ROWS, KARANTINA_WIDTHS,
                  "Sumber: tabel karantina lakehouse. Baris yang nilainya tidak dapat dikonversi tidak dibuang dan tidak pula dipakai — disimpan terpisah agar dapat diperbaiki di sumbernya."),
    "wis": (WIS_HEADER, WIS_ROWS, WIS_WIDTHS,
            "Sumber: serving.mart_wisman (lapisan Gold lakehouse); angka yang sama dipakai dashboard dan dijawab oleh AI Copilot."),
    "dtw": (DTW_HEADER, DTW_ROWS, DTW_WIDTHS,
            "Sumber: serving.mart_kunjungan_dtw — hasil penelusuran sumber publik Juli 2026. Kolom angka sengaja dikosongkan bila tidak ada publikasi resmi; proyeksi disimpan pada kolom terpisah dan tidak pernah dicampur dengan realisasi."),
}

# ---------- shared narrative ----------

PENDAHULUAN_UMUM = (
    "Laporan ini merupakan laporan bulanan ketiga atas pelaksanaan kegiatan Penyediaan dan Pengelolaan "
    "Data Statistik Pariwisata Jakarta Tahun Anggaran 2026. Bila bulan pertama difokuskan pada fondasi data "
    "dan dashboard percontohan, dan bulan kedua pada operasionalisasi platform secara mandiri serta perluasan "
    "data, maka bulan ketiga difokuskan pada penataan data di tingkat arsitektur: (1) membangun dan "
    "mengoperasikan Lake House data pariwisata tiga lapis (mentah–bersih–penyaji) yang berjalan otomatis "
    "setiap hari pada server milik Dinas, sebagai tindak lanjut arahan rapat 13 Juli 2026 mengenai kebutuhan "
    "lake house yang tetap fleksibel bila kelak dipindahkan ke basis data lain; (2) menyatukan lake house "
    "Dinas dengan Rantai Lake — platform lakehouse siap pakai milik penyedia — sebagai alat pengelolaan "
    "sekaligus alat dashboarding, sehingga dashboard dapat disusun tanpa menulis kueri, baik lewat "
    "antarmuka maupun percakapan, dengan jawaban yang selalu menyertakan kueri dan sumber datanya; serta "
    "(3) memastikan seluruh tampilan "
    "data pada platform Dinas dibaca dari lake house tersebut, bukan lagi menembak layanan luar setiap kali "
    "halaman dibuka, sehingga tampilan tetap hidup meskipun layanan sumber sedang terganggu."
)

RTL_UMUM = (
    "Sebagai kelanjutan, tim melanjutkan pengembangan Jakarta Tourism Intelligence Platform — satu fondasi "
    "AI-Data yang menggabungkan lima kemampuan inti: Data Gathering, Lakehouse Making, Dashboarding, AI "
    "Analyst, dan Policy Insight Layer. Pada bulan ketiga ini kemampuan Lakehouse Making telah terwujud dan "
    "berjalan otomatis, sedangkan kemampuan Dashboarding dan AI Analyst dipenuhi dengan menyatukan lake "
    "house Dinas dengan platform Rantai Lake yang sudah tersedia — sehingga anggaran kegiatan tidak dipakai "
    "membangun ulang perangkat yang sudah ada, melainkan untuk data dan penerapannya. Tahap berikutnya "
    "adalah memantapkan lapisan Policy Insight — mengubah jawaban data menjadi "
    "bahan kebijakan yang tetap dapat diaudit sampai ke baris datanya."
)

# ---------- per-member content ----------

MEMBERS = []

# 1. Project Manager
MEMBERS.append(dict(
    filename="01-PM-Analisis-Data-Dashboard",
    jabatan="Project Manager Analisis Data dan Pembangunan Dashboard Pariwisata",
    kualifikasi="S2 / setara, pengalaman minimal 5 tahun",
    pendahuluan_peran=(
        "Sebagai Project Manager Analisis Data dan Pembangunan Dashboard Pariwisata, tenaga ahli bertanggung "
        "jawab merencanakan dan mengoordinasikan seluruh sumber daya tim, menetapkan prioritas pekerjaan bulan "
        "ketiga (pembangunan lake house, penyatuannya dengan platform Rantai Lake, dan pemindahan "
        "seluruh tampilan data ke atas lake house), memantau progres, mengelola risiko, serta memastikan "
        "seluruh keluaran selesai sesuai standar kualitas dan harapan pemilik program."
    ),
    uraian_tugas=[
        "Mengelola dan merencanakan SDM yang dibutuhkan untuk kegiatan Analisis Data dan Pembangunan Dashboard Pariwisata, termasuk analisis, desain, pelaksanaan, pengujian/validasi, dan output yang dihasilkan;",
        "Memberikan laporan kegiatan Analisis Data dan Pembangunan Dashboard Pariwisata berupa progres pekerjaan dan pembaruan;",
        "Memastikan kegiatan selesai tepat waktu, sesuai anggaran, memenuhi standar kualitas, dan harapan pemilik program;",
        "Mengelola risiko dan isu pekerjaan, memastikan penyediaan informasi tepat waktu, serta melakukan mitigasi risiko dan langkah eskalasi.",
    ],
    pelaksanaan_intro=(
        "Pada bulan ketiga, tenaga ahli memimpin penutupan salah satu arahan utama rapat 13 Juli 2026, yaitu "
        "penyediaan lake house data pariwisata. Lake house tiga lapis kini berjalan otomatis setiap hari pada "
        "server Dinas, memuat 225 tabel data mentah, 224 tampilan data bersih, dan 6 mart penyaji dashboard. "
        "Lake house tersebut kemudian disatukan dengan Rantai Lake — platform lakehouse siap pakai milik "
        "penyedia — sebagai alat pengelolaan sekaligus alat dashboarding, tanpa perlu membangun perangkatnya "
        "sendiri. Rincian realisasi terhadap uraian tugas sebagai berikut:"
    ),
    realisasi=[
        ("Perencanaan SDM & prioritas pekerjaan",
         "Menetapkan tiga prioritas bulan ketiga (lake house berjalan otomatis, penyatuan lake house dengan platform Rantai Lake, dan pemindahan seluruh tampilan data ke atas lake house) serta membagi peran tim antara rekayasa data, tata kelola mutu, penyusunan dashboard, dan penyusunan dokumentasi metodologi."),
        ("Pelaporan progres & pembaruan",
         "Menyusun rekapitulasi capaian: lake house 3 lapis berjalan harian (225 tabel mentah, 224 tampilan bersih, 6 mart penyaji); platform Dinas membaca 100% dari lake house; dashboard Wisman dilengkapi pembanding Target RPJMD 2025–2030; pembuat grafik mandiri tersedia di setiap halaman dataset; dokumentasi REST API terbit; buku metodologi statistik pariwisata perkotaan terbit sebagai rujukan."),
        ("Ketepatan waktu, mutu & anggaran",
         "Memastikan seluruh komponen lake house memakai perangkat lunak berlisensi terbuka dan berjalan pada satu server milik Dinas, sehingga tidak menambah biaya langganan; menetapkan gerbang mutu harian sebagai kendali kualitas sebelum data tampil pada dashboard."),
        ("Manajemen risiko & isu",
         "Menetapkan mitigasi atas tiga risiko: (a) ketergantungan pada layanan sumber — data disalin dulu ke lake house sehingga tampilan tetap hidup saat sumber terganggu; (b) keterkuncian teknologi — data mentah disimpan dalam format terbuka sehingga mesin basis data dapat diganti tanpa memindahkan data; (c) risiko angka menyesatkan — angka yang belum terverifikasi tidak ditampilkan sebagai realisasi, melainkan dikarantina atau ditandai."),
    ],
    bukti_intro=(
        "Berikut bukti keluaran kegiatan di bawah koordinasi Project Manager: struktur lake house yang berjalan, "
        "platform data yang membacanya, dan dashboard yang tersusun di atasnya melalui Rantai Lake."
    ),
    bukti=[
        ("lake", "Struktur lake house data pariwisata yang berjalan (26 Agustus 2026):"),
        ("img", "01-lake-catalog.jpg", "Gambar 1. Katalog lake house — 183 dataset primer Satu Data Jakarta, 16 dataset sekunder olahan, 27 model bersih, dan 6 mart penyaji, seluruhnya di server Dinas."),
        ("img", "07-app-home.jpg", "Gambar 2. Platform data Dinas — seluruh angka pada halaman ini dibaca dari lake house, bukan dari pemanggilan layanan luar saat halaman dibuka."),
        ("img", "04-bi-dashboard.jpg", "Gambar 3. Lake house Dinas disatukan dengan platform Rantai Lake — dashboard tersusun dari mart Gold, dapat dibuat lewat antarmuka maupun percakapan."),
    ],
    rtl_peran=(
        "Dalam pengembangan platform, Project Manager berperan memutakhirkan roadmap: Fase 1 (fondasi data & "
        "dashboard) dan Fase 2 (platform mandiri & perluasan data) telah terpenuhi; bulan ini Fase 3 dimulai "
        "dengan berdirinya lake house dan berfungsinya AI Analyst tahap awal. Fokus berikutnya adalah "
        "pemasangan Rantai Lake pada server Dinas untuk dipakai operator internal, serta penyusunan "
        "lapisan Policy Insight yang menautkan setiap angka ke KPI Disparekraf dan indikator daya saing kota."
    ),
    rtl_bukti=[],
    penutup=(
        "Secara keseluruhan, pekerjaan bulan ketiga telah menutup kebutuhan lake house yang diarahkan pada "
        "rapat 13 Juli 2026, memindahkan seluruh tampilan data ke atasnya, dan menambah kemampuan "
        "dashboarding serta asisten data melalui Rantai Lake. Koordinasi tim dilanjutkan untuk pemasangan "
        "Rantai Lake di lingkungan Dinas dan penguatan lapisan analisis kebijakan."
    ),
))

# 2. Database Administrator
MEMBERS.append(dict(
    filename="02-Database-Administrator",
    jabatan="Tenaga Ahli Database Administrator (Senior)",
    kualifikasi="S1 Teknik Informatika/Ilmu Komputer/Sistem Informasi, pengalaman minimal 7 tahun",
    pendahuluan_peran=(
        "Sebagai Database Administrator, tenaga ahli bertanggung jawab menyiapkan dan mengelola basis data "
        "statistik pariwisata pada infrastruktur mandiri Dinas — pada bulan ini berkembang menjadi lake house "
        "tiga lapis — menyusun struktur penyimpanan, menjaga kinerja kueri, serta menetapkan mekanisme "
        "pencadangan, pemulihan, dan pemeliharaan."
    ),
    uraian_tugas=[
        "Merancang struktur basis data statistik pariwisata pada platform Database (logical & physical data model);",
        "Menyusun skema tabel, indeks, constraint, view, dan partisi sesuai kebutuhan analitik;",
        "Melaksanakan loading data hasil pengumpulan, pembersihan, dan validasi ke dalam Database;",
        "Melakukan optimasi kinerja basis data (query tuning, indexing, statistik, partitioning);",
        "Menyusun tata kelola akses dan keamanan basis data (role, privilege, audit trail);",
        "Menyusun pedoman backup, recovery, dan pemeliharaan basis data;",
        "Mendokumentasikan basis data (ERD, kamus tabel, panduan operasional).",
    ],
    pelaksanaan_intro=(
        "Pada bulan ketiga, tenaga ahli membangun lake house data pariwisata di atas server Dinas: data mentah "
        "disimpan dalam format terbuka pada object storage, dibaca oleh satu mesin analitik yang sekaligus "
        "menyajikan tabel agregat untuk dashboard. Struktur ini menjawab kebutuhan agar data tetap dapat "
        "dipindahkan ke basis data lain di kemudian hari tanpa memindahkan isinya. Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Perancangan struktur data tiga lapis",
         "Merancang dan menerapkan lapisan Bronze (mentah, format terbuka Apache Iceberg, lengkap kolom audit: waktu penarikan, sumber, nomor batch, sidik jari baris), Silver (bersih & bertipe), dan Gold (agregat penyaji dashboard)."),
        ("Skema tabel, view & tabel agregat",
         "Menyusun 225 tabel mentah, 224 view bersih hasil inferensi tipe otomatis dengan ambang konversi 95%, serta 6 tabel mart penyaji; pemutakhiran mart memakai pola tabel-bayangan lalu ditukar secara atomik agar dashboard tidak pernah membaca tabel setengah jadi."),
        ("Loading data terkurasi & tervalidasi",
         "Menjalankan pemuatan harian terjadwal (pukul 02.00) untuk seluruh dataset Satu Data Jakarta dan berkas pendataan lapangan; seluruh proses idempoten sehingga pengulangan tidak menimbulkan duplikasi."),
        ("Optimasi kinerja",
         "Menyetel mesin analitik agar dashboard membaca tabel agregat (bukan menghitung ulang dari data mentah); hasil pemantauan menunjukkan waktu respons kueri persentil-95 sebesar 210 ms dengan tingkat kesalahan kueri 0%."),
        ("Tata kelola akses & keamanan",
         "Memisahkan pengguna basis data untuk aplikasi (hanya baca, terbatas pada lapisan penyaji) dan untuk operator; seluruh kredensial dipindahkan ke variabel lingkungan container dan tidak lagi tertulis di dalam kode."),
        ("Backup, recovery & pemeliharaan",
         "Menyiapkan pencadangan harian pukul 04.00 (salinan inkremental bucket data ke bucket cadangan, dapat diarahkan ke penyimpanan terpisah), prosedur pemulihan per tanggal, serta pemeliharaan berkala berupa penghapusan snapshot lama (retensi 7 hari) agar metadata dan penyimpanan tidak membengkak."),
        ("Dokumentasi",
         "Menyusun dokumen operasional lake house (cara menjalankan dari nol, lapisan data, jadwal, prosedur pencadangan/pemulihan) yang disimpan bersama kode."),
    ],
    bukti_intro=(
        "Bukti berupa struktur lake house yang berjalan, isi lapisan penyaji, dan hasil pemantauan kinerja."
    ),
    bukti=[
        ("lake", "Struktur penyimpanan lake house per 26 Agustus 2026:"),
        ("mart", "Isi lapisan penyaji (Gold) yang dibaca dashboard:"),
        ("img", "01-lake-catalog.jpg", "Gambar 1. Katalog lake house dengan pembagian ruang nama (primer, sekunder, bersih, penyaji) beserta kepemilikan data."),
        ("img", "03-lake-observability.jpg", "Gambar 2. Pemantauan kinerja — waktu respons kueri persentil-95 210 ms, kesalahan kueri 0%, tanpa insiden."),
    ],
    rtl_peran=(
        "Pada pengembangan platform, DBA berperan memantapkan operasi lake house: menguji pemulihan dari "
        "cadangan secara berkala, memperluas mart penyaji untuk indikator baru, serta menyiapkan pemisahan "
        "ruang data antar-unit agar platform siap dipakai lebih dari satu bidang di lingkungan Dinas."
    ),
    rtl_bukti=[],
    penutup=(
        "Lake house data pariwisata telah berdiri dan berjalan otomatis di server Dinas, lengkap dengan "
        "pencadangan, pemulihan, dan pemeliharaan. Tahap berikutnya adalah uji pemulihan berkala dan "
        "perluasan mart penyaji untuk indikator baru."
    ),
))

# 3. MDM Specialist
MEMBERS.append(dict(
    filename="03-MDM-Specialist",
    jabatan="Tenaga Ahli Master Data Management (MDM) Specialist (Senior)",
    kualifikasi="S1 Sistem Informasi/Teknik Informatika/Ilmu Komputer/Statistik, pengalaman minimal 7 tahun",
    pendahuluan_peran=(
        "Sebagai MDM Specialist, tenaga ahli bertanggung jawab membentuk dan menjaga data induk entitas "
        "pariwisata serta memastikan setiap nilai yang masuk ke lapisan penyaji telah melewati pemeriksaan "
        "mutu — pada bulan ini diwujudkan sebagai gerbang mutu otomatis yang berjalan harian, lengkap dengan "
        "mekanisme karantina bagi baris yang belum layak dipakai."
    ),
    uraian_tugas=[
        "Mengidentifikasi entitas master data pariwisata Jakarta dan menyusun struktur data master;",
        "Menyusun kebijakan dan prosedur MDM (data ownership, data stewardship, alur perubahan, persetujuan, dan audit);",
        "Melaksanakan pembentukan master data melalui konsolidasi, deduplikasi, verifikasi, dan validasi data dari berbagai sumber;",
        "Menyusun mekanisme pemeliharaan dan pemutakhiran master data secara berkala;",
        "Memastikan integrasi master data ke dalam database dan ketersediaannya bagi Dashboard melalui metadata.",
    ],
    pelaksanaan_intro=(
        "Pada bulan ketiga, pekerjaan MDM berpindah dari perbaikan manual per dataset menjadi aturan mutu yang "
        "dijalankan mesin setiap hari: setiap kolom diperiksa, baris yang nilainya tidak dapat dibaca sebagai "
        "angka atau tanggal dikarantina, dan penurunan jumlah baris yang mencurigakan ditandai sebagai anomali. "
        "Dimensi bersama (negara asal, bulan) dibakukan agar penamaan seragam lintas dataset. Realisasi terhadap "
        "uraian tugas:"
    ),
    realisasi=[
        ("Struktur master & dimensi bersama",
         "Membakukan dimensi bersama pada lake house: dimensi negara asal (menyatukan 61 varian penulisan menjadi 23 nama kanonik) dan dimensi bulan (menyeragamkan penulisan nama bulan), sehingga penggabungan antar-dataset tidak lagi pecah karena beda ejaan."),
        ("Kebijakan & prosedur stewardship",
         "Menetapkan aturan tegas: nilai yang tidak lolos konversi tidak dibuang dan tidak pula dipakai — dipindahkan ke karantina beserta alasan dan contoh nilainya, agar dapat ditagihkan perbaikannya ke pemilik data."),
        ("Pemeriksaan mutu otomatis harian",
         "Mengoperasikan gerbang mutu di antara lapisan bersih dan lapisan penyaji; pada 26 Agustus 2026 tercatat 223 pemeriksaan lolos, 42 peringatan, dan 83 kegagalan yang seluruh barisnya dikarantina sebelum menyentuh dashboard."),
        ("Deteksi anomali & keterlacakan",
         "Menyimpan riwayat jumlah baris setiap tabel (3.130 catatan) untuk mendeteksi data yang tiba-tiba anjlok lebih dari 50%, serta menyimpan seluruh hasil pemeriksaan agar mutu data dapat ditelusuri per tanggal."),
        ("Ketersediaan master bagi dashboard",
         "Memastikan metadata dataset (nama, penanggung jawab, frekuensi, sumber, kamus kolom) ikut tersimpan di lake house sehingga katalog dan dashboard membaca keterangan yang sama dengan datanya."),
    ],
    bukti_intro=(
        "Bukti berupa hasil gerbang mutu harian dan isi karantina yang menahan baris bermasalah agar tidak "
        "menyesatkan indikator."
    ),
    bukti=[
        ("qual", "Hasil gerbang mutu otomatis (contoh satu hari penuh, 26 Agustus 2026):"),
        ("karantina", "Isi karantina — baris yang ditahan beserta alasannya:"),
        ("img", "02-lake-pipelines.jpg", "Gambar 1. Alur pemutakhiran data terjadwal — pemuatan harian dan pemeliharaan/pencadangan, keduanya berstatus selesai."),
        ("img", "01-lake-catalog.jpg", "Gambar 2. Katalog data induk beserta kepemilikan dan pembagian ruang nama pada lake house."),
    ],
    rtl_peran=(
        "Pada pengembangan platform, MDM Specialist berperan menaikkan mutu dari sekadar terdeteksi menjadi "
        "tertagih: menyusun daftar temuan karantina per pemilik data sebagai bahan surat/nota dinas, serta "
        "menyiapkan silsilah data (lineage) agar setiap angka pada dashboard dapat ditelusuri sampai ke baris "
        "sumbernya."
    ),
    rtl_bukti=[],
    penutup=(
        "Pemeriksaan mutu data kini berjalan otomatis setiap hari dengan mekanisme karantina yang tercatat dan "
        "dapat ditelusuri. Tahap berikutnya adalah menindaklanjuti temuan karantina kepada pemilik data dan "
        "melengkapi silsilah data."
    ),
))

# 4. BI Developer
MEMBERS.append(dict(
    filename="04-BI-Developer-1",
    jabatan="Tenaga Ahli BI Developer 1 (Junior)",
    kualifikasi="S1 Sistem Informasi/Teknik Informatika/Ilmu Komputer/Statistik, pengalaman minimal 5 tahun",
    pendahuluan_peran=(
        "Sebagai BI Developer, tenaga ahli bertanggung jawab membangun dashboard indikator pariwisata dan — "
        "pada bulan ini — menyediakan alat agar pengguna di lingkungan Dinas dapat menyusun grafik dan "
        "dashboard sendiri, baik melalui antarmuka maupun melalui percakapan, tanpa perlu menulis kueri."
    ),
    uraian_tugas=[
        "Mengkonfigurasi koneksi Dashboard ke database melalui semantic layer/metadata MDM;",
        "Membangun dashboard interaktif (overview pariwisata, kunjungan wisatawan, akomodasi, destinasi, kontribusi ekonomi, dan dashboard tematik lainnya);",
        "Menyusun visualisasi yang konsisten dengan standar metadata dan definisi data;",
        "Mengatur akses pengguna, jadwal refresh data, dan pemeliharaan dashboard;",
        "Menyusun panduan pemanfaatan dashboard bagi pengguna di lingkungan Dinas.",
    ],
    pelaksanaan_intro=(
        "Pada bulan ketiga, tenaga ahli menyatukan lake house Dinas dengan Rantai Lake — platform lakehouse "
        "siap pakai milik penyedia — lalu menyusun dashboard indikator di atasnya: dashboard bernama, kartu "
        "grafik dari ragam visualisasi yang tersedia, penataan letak, penyaring lintas kartu, sampai berbagi "
        "tautan, tanpa perlu membangun perangkat BI sendiri. Di sisi platform Dinas, dashboard Wisman "
        "diperkaya dengan pembanding target RPJMD dan setiap halaman dataset dilengkapi pembuat grafik "
        "mandiri. Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Koneksi dashboard ke lapisan penyaji",
         "Menyambungkan Rantai Lake ke mart Gold lake house Dinas dan menyiapkan kredensial hanya-baca; kueri tiap kartu disusun sistem dari daftar kolom yang tervalidasi, sehingga pengguna maupun asisten AI tidak pernah menulis kueri mentah."),
        ("Penyusunan dashboard di atas platform",
         "Menyusun dashboard indikator memakai kanvas dan ragam visualisasi yang sudah disediakan platform (grafik, angka kunci/KPI, tabel, teks, peta wilayah), termasuk menata letak kartu dan mode ubah tampilan — tanpa pengembangan perangkat BI dalam kegiatan ini."),
        ("Dashboard tematik & indikator",
         "Menambahkan pembanding Target vs Realisasi Wisman (RPJMD 2025–2030, indikator Jumlah Tamu Mancanegara) pada dashboard Wisatawan Internasional, serta melengkapi tampilan dengan komposisi triwulan, peringkat negara asal, dan peta komposisi (treemap)."),
        ("Pembuat grafik mandiri per dataset",
         "Menyediakan panel Visualisasi Data pada setiap halaman dataset katalog: pengguna memilih tipe grafik, kolom kategori, jenis agregasi, dan kolom nilai, lalu menambahkan grafik — tanpa menulis rumus atau kueri."),
        ("Penyaring, berbagi & pemeliharaan",
         "Mengatur penyaring tahun dan penyaring lintas kartu, penelusuran ke baris mentah saat kartu diklik, ekspor PDF, tautan berbagi yang dapat dicabut, serta penyematan (embed) untuk situs lain — seluruhnya hanya-baca dan mengambil data dari mart penyaji."),
        ("Kewaspadaan & ringkasan berkala",
         "Menyiapkan aturan kewaspadaan (alert) dan ringkasan berkala (digest) atas indikator terpilih: angka dibandingkan terhadap ambang tertentu dan dikirim melalui surel atau kanal pesan."),
    ],
    bukti_intro=(
        "Bukti berupa tangkapan layar dashboard yang disusun di atas Rantai Lake dan hasilnya pada platform Dinas."
    ),
    bukti=[
        ("img", "04-bi-dashboard.jpg", "Gambar 1. Dashboard tersusun dari mart Gold — angka kunci, tren bulanan, negara asal, dan komposisi wilayah, dengan penyaring dan mode ubah tata letak."),
        ("img", "05-bi-chart-gallery.jpg", "Gambar 2. Ragam visualisasi yang tersedia pada platform dan dipakai menyusun dashboard indikator."),
        ("img", "08-wisman-target.jpg", "Gambar 3. Dashboard Wisatawan Internasional — perbandingan Realisasi (jingga) terhadap Target RPJMD 2025–2030 (hijau)."),
        ("img", "09-chart-builder.jpg", "Gambar 4. Pembuat grafik mandiri pada halaman dataset katalog — pilih tipe, kategori, agregasi, dan nilai, lalu tambahkan grafik."),
    ],
    rtl_peran=(
        "Pada pengembangan platform, BI Developer berperan menyusun dashboard baku untuk pimpinan (executive "
        "dashboard) di atas platform yang sudah tersedia, menambah indikator strategis (destinasi, MICE, okupansi "
        "hotel, efektivitas promosi), serta menyiapkan panduan singkat agar staf Dinas dapat menyusun dashboard "
        "sendiri."
    ),
    rtl_bukti=[
        ("06-ai-copilot.jpg", "Gambar 5. Penyusunan dashboard melalui percakapan — pertanyaan pengguna diterjemahkan menjadi kueri ke mart Gold, jawabannya menyertakan kueri dan tabel hasilnya."),
    ],
    penutup=(
        "Dashboard indikator telah tersusun di atas platform Rantai Lake, diperkaya dengan pembanding target "
        "RPJMD, dan setiap dataset kini dapat divisualisasikan langsung oleh pengguna. Tahap berikutnya adalah "
        "menyusun dashboard baku untuk pimpinan dan panduan pemakaiannya."
    ),
))

# 5. Data Engineer
MEMBERS.append(dict(
    filename="05-Data-Engineer-2",
    jabatan="Tenaga Ahli Data Engineer 2 (Junior)",
    kualifikasi="S1 Teknik Informatika/Ilmu Komputer/Sistem Informasi, pengalaman minimal 6 tahun",
    pendahuluan_peran=(
        "Sebagai Data Engineer, tenaga ahli bertanggung jawab merancang dan menjalankan alur pengumpulan dan "
        "pengolahan data — pada bulan ini ditingkatkan dari skrip yang dijalankan manual menjadi alur "
        "terjadwal yang berjalan sendiri setiap hari di atas lake house, lengkap dengan pemantauan dan "
        "pemulihan bila terjadi kegagalan."
    ),
    uraian_tugas=[
        "Merancang dan melaksanakan pipeline pengumpulan dan integrasi data (ETL/ELT) dari berbagai sumber ke dalam Database;",
        "Membersihkan, mentransformasi, dan menstandarisasi raw data agar siap digunakan untuk analisis dan visualisasi;",
        "Memastikan integrasi data multi-sumber dilakukan dengan menjaga kualitas dan keterlacakan data;",
        "Menerapkan standar metadata dan master data dalam proses integrasi data;",
        "Menjamin keamanan data dan kepatuhan terhadap tata kelola data pada proses integrasi.",
    ],
    pelaksanaan_intro=(
        "Pada bulan ketiga, tenaga ahli membangun alur data lake house dari hulu ke hilir: penarikan data ke "
        "lapisan mentah, pembersihan otomatis ke lapisan bersih, dan penyusunan mart penyaji — seluruhnya "
        "diorkestrasi dan dijadwalkan, ditambah alur pemeliharaan dan pencadangan. Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Alur pengumpulan & integrasi terjadwal",
         "Membangun dua alur terjadwal: pemutakhiran lake house setiap hari pukul 02.00 (tarik dataset Satu Data Jakarta + berkas pendataan lapangan → lapisan mentah → bersih → penyaji) dan alur pemeliharaan/pencadangan pukul 04.00; keduanya dapat dipicu manual dan statusnya terpantau."),
        ("Pembersihan & standardisasi otomatis",
         "Menerapkan inferensi tipe otomatis pada 224 tabel: kolom diuji apakah dapat dibaca sebagai angka atau tanggal dengan ambang keberhasilan 95%, termasuk validasi pulang-pergi untuk menolak tanggal mustahil (mis. 31 Februari) dan nilai sentinel."),
        ("Keterlacakan lintas sumber",
         "Menyertakan kolom audit pada setiap baris lapisan mentah (waktu penarikan, alamat sumber, nomor batch, sidik jari baris) dan menyimpan metadata dataset di lake house, sehingga asal setiap angka dapat ditelusuri."),
        ("Penerapan standar metadata & master",
         "Menerapkan dimensi bersama (negara, bulan) pada proses pengolahan dan memastikan katalog beserta kamus kolom ikut termuat, sehingga aplikasi tidak perlu lagi memanggil layanan luar untuk menampilkan keterangan dataset."),
        ("Keamanan & kepatuhan",
         "Memisahkan pengguna basis data untuk aplikasi (hanya baca) dan operator, memindahkan seluruh kredensial ke variabel lingkungan, serta memastikan seluruh komponen lake house berlisensi terbuka (Apache 2.0) sehingga aman dipakai dan tidak menimbulkan kewajiban lisensi."),
        ("Pemantauan & pemulihan",
         "Menyiapkan pemantauan kinerja dan kesehatan layanan, pencadangan harian yang dapat dipulihkan per tanggal, serta pemeliharaan berkala untuk membersihkan snapshot lama."),
    ],
    bukti_intro=(
        "Bukti berupa alur terjadwal yang berjalan beserta hasil pemantauannya, dan struktur data yang "
        "dihasilkannya."
    ),
    bukti=[
        ("img", "02-lake-pipelines.jpg", "Gambar 1. Dua alur data terjadwal (pemutakhiran harian dan pemeliharaan/pencadangan) beserta sumber, sasaran, dan jadwalnya — keduanya berstatus selesai."),
        ("lake", "Keluaran alur data — struktur lake house per 26 Agustus 2026:"),
        ("img", "03-lake-observability.jpg", "Gambar 2. Pemantauan operasional — waktu respons kueri, tingkat kesalahan, keterlambatan pemuatan, dan insiden."),
    ],
    rtl_peran=(
        "Pada pengembangan platform, Data Engineer berperan menambah sumber data baru ke dalam alur yang sudah "
        "berjalan (tanpa membangun ulang), menyiapkan pemuatan bertahap agar tidak menarik ulang seluruh data "
        "setiap hari, serta menyiapkan pemindahan cadangan ke penyimpanan terpisah dari server utama."
    ),
    rtl_bukti=[],
    penutup=(
        "Alur data lake house telah berjalan otomatis setiap hari, terpantau, tercadangkan, dan terpelihara. "
        "Tahap berikutnya adalah pemuatan bertahap dan penambahan sumber data baru ke dalam alur yang sama."
    ),
))

# 6. Data Analyst
MEMBERS.append(dict(
    filename="06-Data-Analyst",
    jabatan="Tenaga Ahli Data Analyst (Intermediate)",
    kualifikasi="S1 Sistem Informasi/Teknik Informatika/Ilmu Komputer/Statistik, pengalaman minimal 6 tahun",
    pendahuluan_peran=(
        "Sebagai Data Analyst, tenaga ahli bertanggung jawab menyiapkan dataset untuk pelaporan dan "
        "visualisasi, menyusun rekap indikator, serta memastikan keakuratan angka yang disajikan — termasuk "
        "menolak angka yang tidak dapat diverifikasi meskipun beredar di ruang publik."
    ),
    uraian_tugas=[
        "Menyiapkan dataset untuk pelaporan dan visualisasi sesuai kebutuhan Dinas;",
        "Menyusun laporan rutin (bulanan, triwulanan, tahunan) dan laporan ad-hoc yang relevan;",
        "Memastikan kualitas dan keakuratan data yang disajikan pada antarmuka pelaporan dan dashboard;",
        "Berkolaborasi dengan pengguna data untuk memvalidasi kebutuhan informasi.",
    ],
    pelaksanaan_intro=(
        "Pada bulan ketiga, tenaga ahli menyusun rekap indikator di atas lapisan penyaji lake house, "
        "menyelesaikan penelusuran data kunjungan 31 daya tarik wisata (DTW) untuk periode Juli 2026, serta "
        "menandai satu angka yang beredar luas namun keliru agar tidak terpakai dalam laporan resmi. "
        "Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Penyiapan dataset untuk pelaporan",
         "Menyusun mart penyaji sebagai sumber tunggal angka laporan: wisman per bulan/negara/kawasan (276 baris), kunjungan DTW beserta sumbernya (35 baris), kesiapan indikator GCI/GPCI (28 baris), event, kuliner, dan pendataan lapangan Atlas."),
        ("Penelusuran data kunjungan 31 DTW",
         "Melakukan dua putaran penyisiran sumber publik untuk kunjungan Juli 2026: dari 31 destinasi, hanya 2 yang memiliki angka terbit — dan keduanya bukan angka bulan penuh. Temuan pokoknya: angka kunjungan hanya dirilis Pemprov saat libur nasional besar, sehingga menunggu publikasi tidak akan pernah menyelesaikan masalah; sumber wajib tetap pengelola destinasi, dengan BPS DKI sebagai rujukan provinsi (jeda dua bulan)."),
        ("Penjaminan akurasi & penolakan angka keliru",
         "Menandai angka “kunjungan Jakarta Semester I 2026 sebesar 9,86 juta” yang beredar di media sebagai tidak konsisten — angka tersebut setara satu bulan wisnus menurut BPS DKI; bila dipakai, kunjungan Jakarta akan tercatat jauh lebih kecil dari kenyataan. Rekomendasi: gunakan angka BPS dengan periode yang eksplisit."),
        ("Pemisahan realisasi dan proyeksi",
         "Menyusun lembar kerja kunjungan DTW dengan kolom realisasi sengaja dikosongkan bila tidak ada publikasi, dan proyeksi ditempatkan pada kolom terpisah sebagai alat kontrol — tidak pernah dicampur menjadi satu angka."),
        ("Rekap indikator & validasi kebutuhan",
         "Menyusun rekap wisatawan mancanegara per kawasan asal dan per negara di atas mart Gold, serta menyelaraskan definisi indikator (mis. Jumlah Tamu Mancanegara pada RPJMD 2025–2030) bersama Business Analyst agar angka dashboard sebanding dengan targetnya."),
    ],
    bukti_intro=(
        "Bukti berupa rekap indikator dari lapisan penyaji dan hasil penelusuran data kunjungan DTW."
    ),
    bukti=[
        ("wis", "Rekap wisatawan mancanegara per kawasan asal (dari mart penyaji lake house):"),
        ("img", "08-wisman-target.jpg", "Gambar 1. Realisasi wisman dibandingkan target RPJMD 2025–2030 pada dashboard indikator."),
        ("dtw", "Hasil penelusuran kunjungan 31 DTW untuk Juli 2026:"),
        ("img", "12-buku-bab.jpg", "Gambar 2. Dokumentasi metodologi — bab normalisasi data kuantitatif & kualitatif pada buku Statistika Pariwisata Perkotaan, dipakai sebagai rujukan penyusunan indikator."),
    ],
    rtl_peran=(
        "Pada pengembangan platform, Data Analyst berperan membakukan laporan rutin di atas mart penyaji dan "
        "memanfaatkan asisten data untuk mempercepat penyusunan laporan ad-hoc, dengan syarat setiap angka "
        "tetap disertai kueri dan sumbernya sehingga dapat diperiksa ulang."
    ),
    rtl_bukti=[],
    penutup=(
        "Rekap indikator kini bersumber tunggal dari lapisan penyaji lake house, dan penelusuran data "
        "kunjungan DTW telah menghasilkan kesimpulan jalur perolehan data yang jelas. Tahap berikutnya adalah "
        "membakukan laporan bulanan dan menempuh jalur permintaan data resmi ke pengelola destinasi."
    ),
))

# 7. Business Analyst
MEMBERS.append(dict(
    filename="07-Business-Analyst-1",
    jabatan="Tenaga Ahli Business Analyst 1 (Intermediate)",
    kualifikasi="S1 Sistem Informasi/Teknik Informatika/Ilmu Komputer/Manajemen/Statistik, pengalaman minimal 6 tahun",
    pendahuluan_peran=(
        "Sebagai Business Analyst, tenaga ahli bertanggung jawab menerjemahkan kebutuhan kebijakan menjadi "
        "kebutuhan data dan arsitektur yang konkret, menyusun gap analysis, mendokumentasikan requirement, "
        "serta memastikan keluaran teknis bulan ini benar-benar menjawab arahan pemangku kepentingan pada "
        "rapat 13 Juli 2026."
    ),
    uraian_tugas=[
        "Melakukan pemetaan kondisi tata kelola data pariwisata saat ini (as-is) melalui wawancara, observasi, dan analisis dokumen;",
        "Merumuskan kondisi tata kelola data yang dituju (to-be) bersama pemangku kepentingan;",
        "Menyusun gap analysis as-is vs to-be beserta rekomendasi pemenuhan;",
        "Menerjemahkan kebutuhan bisnis menjadi spesifikasi tata kelola data (standar data, metadata, kebijakan MDM, kebijakan akses);",
        "Mendokumentasikan requirement, alur proses, dan use case pengelolaan data pariwisata;",
        "Berkoordinasi dengan pemangku kepentingan internal dan eksternal Dinas;",
        "Mendukung penyusunan roadmap dan rencana implementasi standarisasi data.",
    ],
    pelaksanaan_intro=(
        "Pada bulan ketiga, tenaga ahli menerjemahkan dua arahan rapat 13 Juli 2026 — kebutuhan lake house dan "
        "keharusan tetap fleksibel bila kelak dipindahkan ke basis data lain — menjadi spesifikasi arsitektur "
        "yang kemudian dibangun; menyusun dokumentasi antarmuka data (REST API) agar unit lain dapat memakai "
        "data Dinas; serta menyusun buku metodologi sebagai rujukan bersama. Realisasi terhadap uraian tugas:"
    ),
    realisasi=[
        ("Spesifikasi kebutuhan dari arahan rapat",
         "Menerjemahkan arahan lake house menjadi spesifikasi teknis: data disimpan dalam format terbuka di penyimpanan objek, bukan di dalam mesin basis data — sehingga mesin pengolah dapat diganti (termasuk ke basis data komersial) tanpa memindahkan isi data."),
        ("Gap analysis & rekomendasi",
         "Mengidentifikasi kesenjangan as-is → to-be: tampilan data yang bergantung pada layanan luar saat halaman dibuka, pemeriksaan mutu yang manual, dan angka yang tersebar di banyak berkas. Rekomendasi yang telah dipenuhi bulan ini: satu sumber angka (mart penyaji), pemeriksaan mutu otomatis, dan penyajian mandiri."),
        ("Kajian kepatuhan lisensi",
         "Menyusun kajian lisensi seluruh komponen yang dipakai dan menolak komponen yang berisiko bagi Pemerintah Daerah maupun penyediaan layanan; seluruh komponen lake house yang dipakai berlisensi terbuka Apache 2.0."),
        ("Dokumentasi requirement & antarmuka data",
         "Menyusun dan menerbitkan dokumentasi REST API platform (katalog dataset, detail & baris data, ekspor CSV/XLSX, sinkronisasi) sehingga unit lain di lingkungan Pemprov dapat memakai data Dinas secara langsung dan konsisten."),
        ("Dokumentasi metodologi (buku)",
         "Menyusun buku “Statistika Pariwisata Perkotaan” — panduan metodologis dari pengumpulan data sampai peringkat kota global (9 bab, 4 bagian, dilengkapi glosarium, akronim, indeks, etika & regulasi data, serta daftar pustaka) sebagai rujukan bersama Dinas dan mitra."),
        ("Koordinasi & roadmap",
         "Menyelaraskan definisi indikator dengan dokumen perencanaan daerah (RPJMD 2025–2030, indikator Jumlah Tamu Mancanegara) sehingga dashboard menampilkan realisasi berdampingan dengan target resmi, dan memutakhirkan roadmap standardisasi data bersama Project Manager."),
    ],
    bukti_intro=(
        "Bukti berupa dokumentasi antarmuka data yang terbit, buku metodologi, dan perwujudan lapisan analisis "
        "data yang dapat ditanya langsung."
    ),
    bukti=[
        ("img", "10-api-docs.jpg", "Gambar 1. Dokumentasi REST API platform data — endpoint publik tanpa autentikasi untuk katalog, detail data, dan ekspor."),
        ("img", "11-buku-home.jpg", "Gambar 2. Buku “Statistika Pariwisata Perkotaan” — dokumentasi metodologi statistik pariwisata kota, terbit sebagai situs yang dapat ditelusuri."),
        ("img", "06-ai-copilot.jpg", "Gambar 3. Perwujudan lapisan AI Analyst — pertanyaan dalam bahasa sehari-hari dijawab dengan kueri yang ditampilkan terbuka beserta tabel hasilnya dari data Dinas."),
    ],
    rtl_peran=(
        "Pada pengembangan platform, Business Analyst berperan menyusun lapisan Policy Insight: menautkan "
        "setiap indikator ke KPI Disparekraf dan target RPJMD, menyiapkan policy brief berbasis angka yang "
        "dapat diaudit, serta menyusun tata kelola akses data antar-unit di lingkungan Pemprov."
    ),
    rtl_bukti=[],
    penutup=(
        "Arahan rapat 13 Juli 2026 mengenai lake house telah diterjemahkan menjadi spesifikasi dan terbukti "
        "terbangun, dokumentasi antarmuka data serta buku metodologi telah terbit. Tahap berikutnya adalah "
        "penyusunan lapisan analisis kebijakan dan tata kelola akses data antar-unit."
    ),
))

# ---------- laporan utama (payung) ----------

def kop(doc, judul, subjudul):
    para(doc, "PEMERINTAH PROVINSI DAERAH KHUSUS IBUKOTA JAKARTA", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, "DINAS PARIWISATA DAN EKONOMI KREATIF", size=13, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, "Bidang Data, Informasi dan Pengembangan Destinasi", size=10, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=2)
    hr = doc.add_paragraph(); hr.paragraph_format.space_after = Pt(10)
    pPr = hr._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '1a1a1a'); pbdr.append(bottom); pPr.append(pbdr)
    para(doc, judul, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    para(doc, subjudul, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=10)


CAPAIAN_HEADER = ["Keluaran bulan ini", "Wujud nyata", "Status"]
CAPAIAN_ROWS = [
    ["Lake House data pariwisata (Bronze–Silver–Gold)",
     "225 tabel mentah (Iceberg), 224 model bersih, 6 mart penyaji; pemuatan harian 02.00, pemeliharaan & cadangan 04.00",
     "Berjalan otomatis"],
    ["Gerbang mutu data harian + karantina",
     "348 pemeriksaan/hari; baris gagal konversi ditahan di karantina, riwayat jumlah baris tersimpan",
     "Berjalan otomatis"],
    ["Penyatuan dengan platform Rantai Lake",
     "Lake house Dinas tersambung ke platform siap pakai: pengelolaan (katalog, alur data, pemantauan), dashboarding, dan asisten data — perangkatnya tidak dibangun dalam kegiatan ini",
     "Tersambung, uji pakai"],
    ["Aplikasi data Dinas membaca lake house",
     "Katalog & seluruh isi dataset dibaca dari lake house; tidak lagi menembak layanan luar saat halaman dibuka",
     "Berjalan"],
    ["Dashboard indikator diperkaya",
     "Wisman: Target RPJMD 2025–2030 vs realisasi, komposisi triwulan, peringkat negara asal",
     "Berjalan"],
    ["Pembuat grafik mandiri per dataset",
     "Panel Visualisasi Data di tiap halaman dataset katalog (tipe grafik, kategori, agregasi, nilai)",
     "Berjalan"],
    ["Label lapisan medallion di katalog",
     "Kolom Lapisan (Bronze/Silver/Gold) + penyaring & keterangan; dihitung dari keadaan lake house",
     "Selesai dibangun, menunggu penerapan"],
    ["Dokumentasi REST API publik",
     "Katalog, detail & baris data, ekspor CSV/XLSX, sinkronisasi",
     "Terbit"],
    ["Buku “Statistika Pariwisata Perkotaan”",
     "9 bab / 4 bagian + glosarium, akronim, indeks, etika & regulasi data, daftar pustaka",
     "Terbit"],
    ["Penelusuran data kunjungan 31 DTW",
     "Dua putaran penyisiran sumber publik Juli 2026 + jalur perolehan data & kontak pengelola",
     "Selesai"],
]
CAPAIAN_WIDTHS = [4.6, 8.0, 3.0]

KENDALA_HEADER = ["Kendala / risiko", "Dampak", "Tindakan & mitigasi"]
KENDALA_ROWS = [
    ["Kunjungan bulanan 31 DTW tidak terbit",
     "Dari 31 destinasi, hanya 2 punya angka terbit untuk Juli 2026 — keduanya bukan bulan penuh",
     "Angka tidak dikarang: kolom dikosongkan, proyeksi dipisah sebagai alat kontrol. Ditempuh jalur permintaan resmi ke pengelola (17 UPT/UP Pemprov bisa lewat nota dinas)"],
    ["Angka publik yang tidak konsisten",
     "Angka “9,86 juta kunjungan Semester I 2026” yang beredar setara satu bulan wisnus menurut BPS DKI",
     "Ditandai agar tidak dipakai dalam laporan resmi; rujukan yang dipakai adalah BPS dengan periode eksplisit"],
    ["Mutu data sumber",
     "Sebagian kolom berisi nilai yang tidak dapat dibaca sebagai angka (mis. 52.735 nilai pada kolom jumlah)",
     "Ditahan di karantina beserta alasannya — tidak dibuang dan tidak dipakai — sebagai bahan perbaikan ke pemilik data"],
    ["Ketergantungan pada satu server",
     "Seluruh lake house berjalan pada satu server Dinas",
     "Pencadangan harian yang dapat dipulihkan per tanggal; data mentah disimpan format terbuka sehingga mesin pengolah dapat diganti tanpa migrasi data"],
]
KENDALA_WIDTHS = [3.8, 5.0, 6.8]

RENCANA_HEADER = ["Rencana bulan berikutnya", "Sasaran"]
RENCANA_ROWS = [
    ["Penerapan label lapisan medallion ke lingkungan publik", "Katalog publik menampilkan posisi tiap dataset di lake house"],
    ["Pemasangan Rantai Lake di server Dinas", "Operator internal dapat memantau & menjalankan alur data sendiri"],
    ["Dashboard baku untuk pimpinan (executive dashboard)", "Satu halaman ringkas indikator utama Disparekraf"],
    ["Permintaan data resmi ke pengelola 31 DTW", "Angka kunjungan bulanan diperoleh dari sumber, bukan dari penyisiran media"],
    ["Uji pemulihan dari cadangan & pemuatan bertahap", "Keandalan operasi lake house teruji, beban pemuatan harian turun"],
    ["Lapisan Policy Insight", "Angka indikator tertaut ke KPI Disparekraf & target RPJMD sebagai bahan kebijakan"],
]
RENCANA_WIDTHS = [7.8, 7.8]

LAMPIRAN_HEADER = ["No.", "Posisi tenaga ahli", "Berkas laporan"]
LAMPIRAN_ROWS = [
    ["1", "Project Manager Analisis Data dan Pembangunan Dashboard Pariwisata", "01-PM-Analisis-Data-Dashboard"],
    ["2", "Database Administrator (Senior)", "02-Database-Administrator"],
    ["3", "Master Data Management (MDM) Specialist (Senior)", "03-MDM-Specialist"],
    ["4", "BI Developer 1 (Junior)", "04-BI-Developer-1"],
    ["5", "Data Engineer 2 (Junior)", "05-Data-Engineer-2"],
    ["6", "Data Analyst (Intermediate)", "06-Data-Analyst"],
    ["7", "Business Analyst 1 (Intermediate)", "07-Business-Analyst-1"],
]
LAMPIRAN_WIDTHS = [1.0, 9.0, 5.6]


def build_utama():
    """Laporan utama (payung) kegiatan bulan ke-3 — merangkum tujuh laporan tenaga ahli."""
    doc = Document()
    st = doc.styles['Normal']; st.font.name = "Arial"; st.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.2); section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.2)

    kop(doc,
        "LAPORAN BULANAN KEGIATAN",
        "Penyediaan dan Pengelolaan Data Statistik Pariwisata Jakarta — Tahun Anggaran 2026")

    t = doc.add_table(rows=0, cols=2); set_table_borders(t, color="bfbfbf")
    ident = [
        ("Periode Laporan", "Bulan ke-3 (Agustus 2026)"),
        ("Kegiatan", "Analisis Data dan Pembangunan Dashboard Pariwisata"),
        ("Sub Kegiatan", "Perencanaan Daya Tarik Wisata Provinsi"),
        ("Lokasi", "Provinsi DKI Jakarta"),
        ("Jumlah Tenaga Ahli", "7 (tujuh) orang — laporan per orang terlampir"),
        ("Tanggal Verifikasi Angka", "26 Agustus 2026 (dibaca langsung dari lake house)"),
    ]
    for k, v in ident:
        row = t.add_row().cells
        cell_text(row[0], k, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT); set_cell_bg(row[0], "eeeeee")
        cell_text(row[1], v, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        row[0].width = Cm(4.8); row[1].width = Cm(10.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    heading(doc, "I.  RINGKASAN EKSEKUTIF")
    para(doc, PENDAHULUAN_UMUM)
    para(doc,
         "Inti capaian bulan ini: data pariwisata Dinas tidak lagi berpindah-pindah antar berkas dan "
         "layanan, melainkan mengalir pada satu jalur yang sama setiap hari — ditarik ke lake house, "
         "dibersihkan, diperiksa mutunya, lalu disajikan ke dashboard. Karena jalur itu berjalan sendiri "
         "dan tercatat, setiap angka yang tampil dapat ditelusuri kembali sampai ke sumber dan tanggal "
         "penarikannya; sebaliknya, angka yang belum layak tidak ikut tampil, melainkan ditahan di "
         "karantina untuk diperbaiki di sumbernya.")

    heading(doc, "II.  CAPAIAN BULAN INI")
    para(doc, "Seluruh keluaran di bawah ini dapat diperiksa langsung pada sistem yang berjalan:",
         space_after=4)
    data_table(doc, CAPAIAN_HEADER, CAPAIAN_ROWS, CAPAIAN_WIDTHS)

    heading(doc, "III.  STRUKTUR DATA YANG DIBANGUN")
    para(doc,
         "Data disusun berlapis (arsitektur medallion). Lapisan mentah disimpan dalam format terbuka di "
         "penyimpanan objek — bukan di dalam mesin basis data — sehingga mesin pengolah dapat diganti di "
         "kemudian hari, termasuk ke basis data komersial, tanpa memindahkan isi datanya. Ini memenuhi "
         "arahan rapat 13 Juli 2026.", space_after=6)
    data_table(doc, LAKE_HEADER, LAKE_ROWS, LAKE_WIDTHS)
    para(doc, "Lapisan penyaji yang dibaca dashboard:", bold=True, size=10.5,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=3)
    data_table(doc, MART_HEADER, MART_ROWS, MART_WIDTHS)

    heading(doc, "IV.  PENJAMINAN MUTU DATA")
    para(doc,
         "Antara lapisan bersih dan lapisan penyaji dipasang gerbang mutu yang berjalan otomatis setiap "
         "hari. Hasilnya tersimpan sehingga mutu data dapat ditelusuri per tanggal.", space_after=6)
    data_table(doc, QUAL_HEADER, QUAL_ROWS, QUAL_WIDTHS)
    data_table(doc, KARANTINA_HEADER, KARANTINA_ROWS, KARANTINA_WIDTHS)

    heading(doc, "V.  BUKTI PELAKSANAAN")
    para(doc,
         "Tangkapan layar berikut diambil dari sistem yang sedang berjalan pada 26 Agustus 2026.",
         space_after=6)
    for f, c in [
        ("01-lake-catalog.jpg", "Gambar 1. Katalog lake house — 183 dataset primer Satu Data Jakarta, 16 dataset sekunder olahan, 27 model bersih, dan 6 mart penyaji."),
        ("02-lake-pipelines.jpg", "Gambar 2. Dua alur data terjadwal (pemutakhiran harian & pemeliharaan/pencadangan) — keduanya berstatus selesai."),
        ("03-lake-observability.jpg", "Gambar 3. Pemantauan operasional — waktu respons kueri persentil-95 210 ms, kesalahan kueri 0%, tanpa insiden."),
        ("04-bi-dashboard.jpg", "Gambar 4. Dashboard indikator yang disusun di atas platform Rantai Lake — kartu membaca mart penyaji lake house Dinas."),
        ("06-ai-copilot.jpg", "Gambar 5. Asisten data (AI Copilot) — pertanyaan bahasa sehari-hari dijawab dengan kueri yang ditampilkan terbuka beserta hasilnya."),
        ("08-wisman-target.jpg", "Gambar 6. Dashboard Wisatawan Internasional — realisasi (jingga) terhadap target RPJMD 2025–2030 (hijau)."),
        ("09-chart-builder.jpg", "Gambar 7. Pembuat grafik mandiri pada halaman dataset katalog."),
        ("10-api-docs.jpg", "Gambar 8. Dokumentasi REST API platform data."),
        ("11-buku-home.jpg", "Gambar 9. Buku “Statistika Pariwisata Perkotaan” sebagai dokumentasi metodologi."),
    ]:
        add_image(doc, f, c)

    heading(doc, "VI.  REKAP INDIKATOR TERPILIH")
    para(doc, "Wisatawan mancanegara per kawasan asal (dibaca dari lapisan penyaji):", bold=True,
         size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=3)
    data_table(doc, WIS_HEADER, WIS_ROWS, WIS_WIDTHS)
    para(doc, "Hasil penelusuran kunjungan 31 daya tarik wisata untuk Juli 2026:", bold=True,
         size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=3)
    data_table(doc, DTW_HEADER, DTW_ROWS, DTW_WIDTHS)

    heading(doc, "VII.  KENDALA, RISIKO, DAN MITIGASI")
    data_table(doc, KENDALA_HEADER, KENDALA_ROWS, KENDALA_WIDTHS)

    heading(doc, "VIII.  RENCANA BULAN BERIKUTNYA")
    para(doc, RTL_UMUM, space_after=6)
    data_table(doc, RENCANA_HEADER, RENCANA_ROWS, RENCANA_WIDTHS)

    heading(doc, "IX.  PENUTUP")
    para(doc,
         "Pekerjaan bulan ketiga menutup kebutuhan lake house yang diarahkan pada rapat 13 Juli 2026, "
         "memindahkan seluruh tampilan data Dinas ke atasnya, dan menambah kemampuan dashboarding serta "
         "asisten data. Seluruh angka pada laporan ini diverifikasi langsung dari sistem yang berjalan "
         "pada 26 Agustus 2026; keluaran yang belum selesai dinyatakan apa adanya beserta rencana "
         "penyelesaiannya.")

    heading(doc, "LAMPIRAN — LAPORAN PELAKSANAAN TUGAS TENAGA AHLI")
    para(doc, "Laporan per tenaga ahli disertakan sebagai satu berkas terpisah untuk masing-masing posisi:",
         space_after=4)
    data_table(doc, LAMPIRAN_HEADER, LAMPIRAN_ROWS, LAMPIRAN_WIDTHS)

    # tanda tangan
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    sig = doc.add_table(rows=0, cols=2)
    r = sig.add_row().cells
    cell_text(r[0], "Mengetahui,\nPejabat Pembuat Komitmen", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(r[1], "Jakarta, ......................... 2026\nProject Manager Kegiatan",
              size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4): sig.add_row()
    r = sig.add_row().cells
    cell_text(r[0], "(  Bima Agung  )", bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(r[1], "(  [Nama Project Manager]  )", bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = sig.add_row().cells
    cell_text(r[0], "NIP. 197907162011011008", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)
    cell_text(r[1], "Project Manager Analisis Data dan Pembangunan Dashboard Pariwisata",
              size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY)

    outdir = os.path.join(BASE, "bulan-3")
    os.makedirs(outdir, exist_ok=True)
    out = os.path.join(outdir, "00-Laporan-Utama-Bulan-3.docx")
    doc.save(out)
    print("wrote", out)


if __name__ == "__main__":
    build_utama()
    for m in MEMBERS:
        build(m)
    print("done: 1 laporan utama +", len(MEMBERS), "laporan tenaga ahli")
